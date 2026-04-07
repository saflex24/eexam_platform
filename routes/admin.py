# ═══════════════════════════════════════════════════════════════════════════
# ADMIN ROUTES - FIXED VERSION
# Fixed:
#   • Random question/option ordering via ExamSession.question_order JSON
#   • result_details_view: correct unanswered count, correct-answer lookup for MCQ
#   • generate_result_pdf: respects per-session question & option order
#   • generate_detailed_report: fixed selected_option / correct_answer access
#   • exam_results: added missing User join for search
#   • log_proctoring_violation: fixed indentation (was unreachable)
#   • All existing logic flow preserved
# ═══════════════════════════════════════════════════════════════════════════

from flask import Blueprint, render_template, request, redirect, url_for, flash, jsonify, send_file, make_response
from flask_login import current_user, login_required
from extensions import db
from models.user import User, Student, Teacher, Role
from models.class_model import Class
from utils.decorators import admin_required
from utils.file_handler import (
    save_upload_file, import_students_from_csv, import_teachers_from_csv,
    generate_student_template, generate_teacher_template
)
from datetime import datetime, timedelta
from collections import defaultdict
from functools import wraps
import pandas as pd
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl import Workbook
import tempfile
import os
import io
import json
import traceback

# Import exam-related models
from models.exam import Exam, ExamResult, ProctoringLog, ExamSession, StudentAnswer, Question, QuestionOption

admin_bp = Blueprint('admin', __name__)


# ══════════════════════════════════════════════════════════════════════════
# HELPER: Resolve question/option order for a session
# ══════════════════════════════════════════════════════════════════════════

def get_ordered_questions(exam_id, exam_session=None):
    """
    Return Question objects in the order the student actually saw them.
    If the ExamSession stored a randomised question_order JSON list of IDs,
    use that; otherwise fall back to Question.order / Question.id.
    """
    base_qs = {q.id: q for q in Question.query.filter_by(exam_id=exam_id).all()}

    if exam_session and hasattr(exam_session, 'question_order') and exam_session.question_order:
        try:
            ordered_ids = json.loads(exam_session.question_order)
            # Return only IDs that still exist
            ordered = [base_qs[qid] for qid in ordered_ids if qid in base_qs]
            # Append any questions not listed (safety net)
            listed = set(ordered_ids)
            extras = sorted(
                [q for q in base_qs.values() if q.id not in listed],
                key=lambda q: (q.order or 0, q.id)
            )
            return ordered + extras
        except (json.JSONDecodeError, KeyError, TypeError):
            pass

    # Default: sort by order then id
    return sorted(base_qs.values(), key=lambda q: (q.order or 0, q.id))


def get_ordered_options(question, exam_session=None):
    """
    Return QuestionOption objects in the order the student actually saw them.
    If ExamSession stored option_order as a JSON dict {str(question_id): [option_ids]},
    use that; otherwise fall back to option_label sort.
    """
    options = list(question.options)

    if exam_session and hasattr(exam_session, 'option_order') and exam_session.option_order:
        try:
            option_order_map = json.loads(exam_session.option_order)
            ordered_ids = option_order_map.get(str(question.id))
            if ordered_ids:
                opt_by_id = {o.id: o for o in options}
                ordered = [opt_by_id[oid] for oid in ordered_ids if oid in opt_by_id]
                listed = set(ordered_ids)
                extras = [o for o in options if o.id not in listed]
                return ordered + extras
        except (json.JSONDecodeError, KeyError, TypeError):
            pass

    return sorted(options, key=lambda o: o.option_label or '')


def get_correct_answer_text(question):
    """Return a human-readable correct answer string for any question type."""
    if question.question_type in ('mcq', 'true_false'):
        correct_opts = [o for o in question.options if getattr(o, 'is_correct', False)]
        if correct_opts:
            o = correct_opts[0]
            return f"{o.option_label}. {o.option_text}"
        return 'N/A'
    # Theory / short-answer: use correct_answer field if it exists
    return getattr(question, 'correct_answer', None) or 'N/A'


def get_student_answer_text(answer, question):
    """Return human-readable student answer string."""
    if not answer:
        return '—'
    if answer.selected_option_id:
        opt = QuestionOption.query.get(answer.selected_option_id)
        if opt:
            return f"{opt.option_label}. {opt.option_text}"
        return 'Unknown option'
    if answer.theory_answer:
        return answer.theory_answer
    return '—'


# ══════════════════════════════════════════════════════════════════════════
# DASHBOARD
# ══════════════════════════════════════════════════════════════════════════

@admin_bp.route('/dashboard')
@admin_required
def dashboard():
    """Admin dashboard with statistics"""
    try:
        print("=== ADMIN DASHBOARD START ===")

        total_students = Student.query.filter_by(is_active=True).count()
        total_teachers = Teacher.query.filter_by(is_active=True).count()
        total_classes  = Class.query.filter_by(is_active=True).count()
        total_exams    = Exam.query.filter_by(is_deleted=False).count()

        print(f"✓ Students: {total_students}, Teachers: {total_teachers}, "
              f"Classes: {total_classes}, Exams: {total_exams}")

        recent_results = ExamResult.query.order_by(ExamResult.submitted_at.desc()).limit(10).all()

        all_results = ExamResult.query.all()
        if all_results:
            avg_score        = sum(r.percentage for r in all_results) / len(all_results)
            total_submissions = len(all_results)
        else:
            avg_score = total_submissions = 0

        top_students = []
        try:
            active_students   = Student.query.filter_by(is_active=True).all()
            student_averages  = []
            for student in active_students:
                student_results = ExamResult.query.filter_by(student_id=student.user_id).all()
                if student_results:
                    avg_pct = sum(r.percentage for r in student_results) / len(student_results)
                    student_averages.append((student, avg_pct))
            student_averages.sort(key=lambda x: x[1], reverse=True)
            top_students = student_averages[:10]
            print(f"✓ Top students: {len(top_students)}")
        except Exception as e:
            print(f"⚠ Error calculating top students: {str(e)}")
            top_students = []

        return render_template('admin/dashboard.html',
                               total_students=total_students,
                               total_teachers=total_teachers,
                               total_classes=total_classes,
                               total_exams=total_exams,
                               avg_score=round(avg_score, 2),
                               total_submissions=total_submissions,
                               recent_results=recent_results,
                               top_students=top_students)

    except Exception as e:
        print(f"=== ERROR IN DASHBOARD ===\n{str(e)}")
        traceback.print_exc()
        flash('Error loading dashboard. Please try again.', 'danger')
        return redirect(url_for('index'))


# ══════════════════════════════════════════════════════════════════════════
# ADVANCED ANALYTICS
# ══════════════════════════════════════════════════════════════════════════

@admin_bp.route('/analytics/advanced')
@admin_required
def advanced_analytics():
    """Advanced analytics dashboard with AI insights"""
    try:
        selected_class   = request.args.get('class_id', type=int)
        selected_subject = request.args.get('subject', '')
        selected_period  = request.args.get('period', 'all')
        selected_exam    = request.args.get('exam_id', type=int)

        results_query = ExamResult.query

        if selected_class:
            students   = Student.query.filter_by(class_id=selected_class, is_active=True).all()
            student_ids = [s.user_id for s in students]
            results_query = (results_query.filter(ExamResult.student_id.in_(student_ids))
                             if student_ids
                             else results_query.filter(ExamResult.student_id == -1))

        if selected_exam:
            results_query = results_query.filter_by(exam_id=selected_exam)
        elif selected_subject:
            subject_exams = Exam.query.filter_by(subject=selected_subject, is_deleted=False).all()
            exam_ids = [e.id for e in subject_exams]
            results_query = (results_query.filter(ExamResult.exam_id.in_(exam_ids))
                             if exam_ids
                             else results_query.filter(ExamResult.exam_id == -1))

        if selected_period != 'all':
            now  = datetime.utcnow()
            days = {'week': 7, 'month': 30, 'quarter': 90, 'year': 365}.get(selected_period, 365)
            results_query = results_query.filter(ExamResult.submitted_at >= now - timedelta(days=days))

        all_results = results_query.all()

        if all_results:
            total_attempts  = len(all_results)
            avg_score       = round(sum(r.percentage for r in all_results) / total_attempts, 2)
            pass_count      = sum(1 for r in all_results if r.is_passed)
            pass_rate       = round((pass_count / total_attempts) * 100, 2)
            unique_students = len(set(r.student_id for r in all_results))
            unique_exams    = len(set(r.exam_id    for r in all_results))
        else:
            avg_score = pass_rate = unique_students = unique_exams = 0

        overall_stats = {
            'total_students': unique_students,
            'avg_score':      avg_score,
            'pass_rate':      pass_rate,
            'total_exams':    unique_exams
        }

        excellent    = sum(1 for r in all_results if r.percentage >= 80)
        average_perf = sum(1 for r in all_results if 50 <= r.percentage < 80)
        poor         = sum(1 for r in all_results if r.percentage < 50)

        performance_dist_labels = ['Excellent (≥80%)', 'Average (50-79%)', 'Poor (<50%)']
        performance_dist_data   = [excellent, average_perf, poor]

        grade_counts = {'A': 0, 'B': 0, 'C': 0, 'D': 0, 'F': 0}
        for r in all_results:
            if   r.percentage >= 90: grade_counts['A'] += 1
            elif r.percentage >= 75: grade_counts['B'] += 1
            elif r.percentage >= 60: grade_counts['C'] += 1
            elif r.percentage >= 50: grade_counts['D'] += 1
            else:                    grade_counts['F'] += 1

        grade_labels = list(grade_counts.keys())
        grade_data   = list(grade_counts.values())

        trend_data_dict = defaultdict(list)
        for r in all_results:
            trend_data_dict[r.submitted_at.strftime('%Y-%m-%d')].append(r.percentage)
        trend_labels = sorted(trend_data_dict.keys())[-10:]
        trend_data   = [round(sum(trend_data_dict[d]) / len(trend_data_dict[d]), 2) for d in trend_labels]

        subject_performance = defaultdict(list)
        for r in all_results:
            exam = Exam.query.get(r.exam_id)
            if exam and exam.subject:
                subject_performance[exam.subject].append(r.percentage)
        subject_labels = list(subject_performance.keys())
        subject_data   = [round(sum(v) / len(v), 2) if v else 0 for v in subject_performance.values()]

        class_performance = {}
        for cls in Class.query.filter_by(is_active=True).all():
            cls_students = Student.query.filter_by(class_id=cls.id, is_active=True).all()
            cls_ids      = [s.user_id for s in cls_students]
            cls_results  = [r for r in all_results if r.student_id in cls_ids]
            if cls_results:
                cls_avg       = sum(r.percentage for r in cls_results) / len(cls_results)
                cls_pass_rate = (sum(1 for r in cls_results if r.is_passed) / len(cls_results)) * 100
                class_performance[cls.name] = {
                    'avg':       round(cls_avg, 2),
                    'pass_rate': round(cls_pass_rate, 2)
                }

        class_labels    = list(class_performance.keys())
        class_avg_data  = [class_performance[c]['avg']       for c in class_labels]
        class_pass_data = [class_performance[c]['pass_rate'] for c in class_labels]

        student_performances = defaultdict(list)
        for r in all_results:
            student_performances[r.student_id].append(r)

        top_performers  = []
        low_performers  = []
        student_insights = []

        for student_id, results in student_performances.items():
            user    = User.query.get(student_id)
            student = Student.query.filter_by(user_id=student_id).first()
            if not user or not student:
                continue

            avg          = sum(r.percentage for r in results) / len(results)
            failed_count = sum(1 for r in results if not r.is_passed)

            if avg >= 80:
                insight        = "High academic achiever with consistent excellent performance."
                concern        = "Risk of burnout or lack of academic challenge."
                recommendation = "Provide advanced materials, competitions, or mentorship roles."
            elif 50 <= avg < 80:
                insight        = "Moderate performance with room for improvement."
                concern        = "Inconsistent scores across exams."
                recommendation = "Provide targeted practice, tutorials, and formative assessments."
            else:
                insight        = "Student is academically struggling."
                concern        = "High failure rate and poor subject comprehension."
                recommendation = "Immediate academic intervention, tutoring, and parental engagement."

            student_insights.append({
                "name":           user.full_name,
                "class_name":     student.class_info.name if student.class_info else "N/A",
                "avg_score":      round(avg, 2),
                "failed_count":   failed_count,
                "insight":        insight,
                "concern":        concern,
                "recommendation": recommendation
            })

            if avg >= 80:
                top_performers.append({
                    'name':       user.full_name,
                    'class_name': student.class_info.name if student.class_info else 'N/A',
                    'avg_score':  round(avg, 2),
                    'exam_count': len(results)
                })
            if avg < 50:
                low_performers.append({
                    'name':         user.full_name,
                    'class_name':   student.class_info.name if student.class_info else 'N/A',
                    'avg_score':    round(avg, 2),
                    'exam_count':   len(results),
                    'failed_count': failed_count
                })

        top_performers.sort(key=lambda x: x['avg_score'], reverse=True)
        low_performers.sort(key=lambda x: x['avg_score'])

        all_classes  = Class.query.filter_by(is_active=True).all()
        all_subjects = Exam.query.filter_by(is_deleted=False).with_entities(Exam.subject).distinct().all()
        subject_list = [s[0] for s in all_subjects if s[0]]
        all_exams    = Exam.query.filter_by(is_deleted=False).all()

        return render_template('admin/advanced_analytics.html',
                               overall_stats=overall_stats,
                               performance_dist_labels=performance_dist_labels,
                               performance_dist_data=performance_dist_data,
                               grade_labels=grade_labels,
                               grade_data=grade_data,
                               trend_labels=trend_labels,
                               trend_data=trend_data,
                               subject_labels=subject_labels,
                               subject_data=subject_data,
                               class_labels=class_labels,
                               class_avg_data=class_avg_data,
                               class_pass_data=class_pass_data,
                               top_performers=top_performers[:10],
                               low_performers=low_performers[:10],
                               student_insights=student_insights,
                               all_classes=all_classes,
                               all_subjects=subject_list,
                               all_exams=all_exams,
                               selected_class=selected_class,
                               selected_subject=selected_subject,
                               selected_period=selected_period,
                               selected_exam=selected_exam)

    except Exception as e:
        print(f"ERROR in advanced analytics: {str(e)}")
        traceback.print_exc()
        flash('Error loading analytics. Please try again.', 'danger')
        return redirect(url_for('admin.dashboard'))


# ══════════════════════════════════════════════════════════════════════════
# EXPORT ANALYTICS
# ══════════════════════════════════════════════════════════════════════════

@admin_bp.route('/analytics/export')
@admin_required
def export_analytics():
    """Export analytics data to Excel"""
    try:
        selected_class   = request.args.get('class_id', type=int)
        selected_subject = request.args.get('subject', '')
        selected_period  = request.args.get('period', 'all')

        results_query = ExamResult.query

        if selected_class:
            students    = Student.query.filter_by(class_id=selected_class, is_active=True).all()
            student_ids = [s.user_id for s in students]
            if student_ids:
                results_query = results_query.filter(ExamResult.student_id.in_(student_ids))

        if selected_subject:
            subject_exams = Exam.query.filter_by(subject=selected_subject, is_deleted=False).all()
            exam_ids = [e.id for e in subject_exams]
            if exam_ids:
                results_query = results_query.filter(ExamResult.exam_id.in_(exam_ids))

        if selected_period != 'all':
            now  = datetime.utcnow()
            days = {'week': 7, 'month': 30, 'quarter': 90, 'year': 365}.get(selected_period, 365)
            results_query = results_query.filter(ExamResult.submitted_at >= now - timedelta(days=days))

        all_results = results_query.all()

        export_data = []
        for result in all_results:
            student = Student.query.filter_by(user_id=result.student_id).first()
            exam    = Exam.query.get(result.exam_id)
            if student and exam:
                export_data.append({
                    'Student Name':    student.user.full_name,
                    'Admission Number': student.admission_number,
                    'Class':           student.class_info.name if student.class_info else 'N/A',
                    'Exam':            exam.title,
                    'Subject':         exam.subject,
                    'Total Marks':     result.total_marks,
                    'Marks Obtained':  result.marks_obtained,
                    'Percentage':      round(result.percentage, 2),
                    'Status':          'Pass' if result.is_passed else 'Fail',
                    'Submitted At':    result.submitted_at.strftime('%Y-%m-%d %H:%M')
                })

        df     = pd.DataFrame(export_data)
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            df.to_excel(writer, index=False, sheet_name='Analytics')
        output.seek(0)

        filename = f"Analytics_Export_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=filename
        )

    except Exception as e:
        print(f"ERROR exporting analytics: {str(e)}")
        traceback.print_exc()
        flash('Error exporting analytics. Please try again.', 'danger')
        return redirect(url_for('admin.advanced_analytics'))


# ══════════════════════════════════════════════════════════════════════════
# CLASS MANAGEMENT
# ══════════════════════════════════════════════════════════════════════════

@admin_bp.route('/classes', methods=['GET'])
@admin_required
def manage_classes():
    try:
        page   = request.args.get('page', 1, type=int)
        search = request.args.get('search', '', type=str)
        query  = Class.query.filter_by(is_active=True)
        if search:
            query = query.filter(Class.name.ilike(f'%{search}%'))
        classes = query.paginate(page=page, per_page=20, error_out=False)
        return render_template('admin/classes.html', classes=classes, search=search)
    except Exception as e:
        print(f"ERROR in manage_classes: {str(e)}")
        traceback.print_exc()
        flash(f'Error loading classes: {str(e)}', 'danger')
        return redirect(url_for('admin.dashboard'))


@admin_bp.route('/class/create', methods=['GET', 'POST'])
@admin_required
def create_class():
    if request.method == 'POST':
        name          = request.form.get('name', '').strip()
        code          = request.form.get('code', '').strip()
        section       = request.form.get('section', '').strip()
        academic_year = request.form.get('academic_year', '').strip()
        description   = request.form.get('description', '').strip()

        if not name or not code or not academic_year:
            flash('Name, code, and academic year are required.', 'danger')
            return redirect(url_for('admin.create_class'))
        if Class.query.filter_by(code=code).first():
            flash('Class code already exists.', 'danger')
            return redirect(url_for('admin.create_class'))

        try:
            class_obj = Class(name=name, code=code, section=section,
                              academic_year=academic_year, description=description)
            db.session.add(class_obj)
            db.session.commit()
            flash('Class created successfully.', 'success')
            return redirect(url_for('admin.manage_classes'))
        except Exception as e:
            db.session.rollback()
            flash(f'Error creating class: {str(e)}', 'danger')

    return render_template('admin/class_form.html', class_obj=None)


@admin_bp.route('/class/<int:id>/edit', methods=['GET', 'POST'])
@admin_required
def edit_class(id):
    try:
        class_obj = Class.query.get_or_404(id)
        if request.method == 'POST':
            name          = request.form.get('name', '').strip()
            code          = request.form.get('code', '').strip()
            section       = request.form.get('section', '').strip()
            academic_year = request.form.get('academic_year', '').strip()
            description   = request.form.get('description', '').strip()

            if not name or not code or not academic_year:
                flash('Name, code, and academic year are required.', 'danger')
                return redirect(url_for('admin.edit_class', id=id))

            existing = Class.query.filter_by(code=code).first()
            if existing and existing.id != id:
                flash('Class code already exists.', 'danger')
                return redirect(url_for('admin.edit_class', id=id))

            try:
                class_obj.name          = name
                class_obj.code          = code
                class_obj.section       = section or None
                class_obj.academic_year = academic_year
                class_obj.description   = description or None
                db.session.commit()
                flash('Class updated successfully.', 'success')
                return redirect(url_for('admin.manage_classes'))
            except Exception as e:
                db.session.rollback()
                flash(f'Error updating class: {str(e)}', 'danger')
                traceback.print_exc()

        return render_template('admin/class_form.html', class_obj=class_obj)
    except Exception as e:
        print(f"ERROR in edit_class: {str(e)}")
        traceback.print_exc()
        flash(f'Error loading class: {str(e)}', 'danger')
        return redirect(url_for('admin.manage_classes'))


@admin_bp.route('/class/<int:id>/delete', methods=['POST'])
@admin_required
def delete_class(id):
    class_obj = Class.query.get_or_404(id)
    try:
        class_obj.is_active = False
        db.session.commit()
        flash('Class deleted successfully.', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Error deleting class: {str(e)}', 'danger')
    return redirect(url_for('admin.manage_classes'))


# ══════════════════════════════════════════════════════════════════════════
# ADMIN PROFILE MANAGEMENT
# ══════════════════════════════════════════════════════════════════════════

@admin_bp.route('/profile', methods=['GET'])
@admin_required
def admin_profile():
    try:
        admin_user = current_user
        stats = {
            'total_students':  Student.query.filter_by(is_active=True).count(),
            'total_teachers':  Teacher.query.filter_by(is_active=True).count(),
            'total_exams':     Exam.query.filter_by(is_deleted=False).count(),
            'total_classes':   Class.query.filter_by(is_active=True).count(),
            'account_created': admin_user.created_at,
            'last_updated':    admin_user.updated_at
        }
        return render_template('admin/profile.html', admin=admin_user, stats=stats)
    except Exception as e:
        print(f"ERROR in admin_profile: {str(e)}")
        traceback.print_exc()
        flash('Error loading profile.', 'danger')
        return redirect(url_for('admin.dashboard'))


@admin_bp.route('/profile/edit', methods=['GET', 'POST'])
@admin_required
def edit_admin_profile():
    try:
        admin_user = current_user
        if request.method == 'POST':
            admin_user.first_name = request.form.get('first_name', '').strip()
            admin_user.last_name  = request.form.get('last_name', '').strip()

            new_email = request.form.get('email', '').strip()
            if new_email and new_email != admin_user.email:
                existing = User.query.filter_by(email=new_email).first()
                if existing and existing.id != admin_user.id:
                    flash('Email already exists.', 'danger')
                    return redirect(url_for('admin.edit_admin_profile'))
                admin_user.email = new_email

            new_username = request.form.get('username', '').strip()
            if new_username and new_username != admin_user.username:
                existing = User.query.filter_by(username=new_username).first()
                if existing and existing.id != admin_user.id:
                    flash('Username already exists.', 'danger')
                    return redirect(url_for('admin.edit_admin_profile'))
                admin_user.username = new_username

            admin_user.gender = request.form.get('gender', 'Other')

            if 'profile_picture' in request.files:
                file = request.files['profile_picture']
                if file and file.filename:
                    try:
                        from werkzeug.utils import secure_filename
                        filename    = secure_filename(file.filename)
                        timestamp   = datetime.now().strftime('%Y%m%d_%H%M%S')
                        filename    = f"admin_{admin_user.id}_{timestamp}_{filename}"
                        upload_dir  = os.path.join('static', 'uploads', 'profile_pictures')
                        os.makedirs(upload_dir, exist_ok=True)
                        file.save(os.path.join(upload_dir, filename))
                        admin_user.profile_picture = filename
                    except Exception as e:
                        print(f"Error uploading profile picture: {str(e)}")
                        flash('Error uploading profile picture.', 'warning')

            admin_user.updated_at = datetime.utcnow()
            try:
                db.session.commit()
                flash('Profile updated successfully!', 'success')
                return redirect(url_for('admin.admin_profile'))
            except Exception as e:
                db.session.rollback()
                flash(f'Error updating profile: {str(e)}', 'danger')

        return render_template('admin/edit_profile.html', admin=admin_user)
    except Exception as e:
        print(f"ERROR in edit_admin_profile: {str(e)}")
        traceback.print_exc()
        flash('Error editing profile.', 'danger')
        return redirect(url_for('admin.admin_profile'))


@admin_bp.route('/profile/change-password', methods=['GET', 'POST'])
@admin_required
def change_admin_password():
    try:
        if request.method == 'POST':
            current_password = request.form.get('current_password', '')
            new_password     = request.form.get('new_password', '')
            confirm_password = request.form.get('confirm_password', '')

            if not all([current_password, new_password, confirm_password]):
                flash('All fields are required.', 'danger')
                return redirect(url_for('admin.change_admin_password'))
            if not current_user.check_password(current_password):
                flash('Current password is incorrect.', 'danger')
                return redirect(url_for('admin.change_admin_password'))
            if new_password != confirm_password:
                flash('New passwords do not match.', 'danger')
                return redirect(url_for('admin.change_admin_password'))
            if len(new_password) < 6:
                flash('Password must be at least 6 characters long.', 'danger')
                return redirect(url_for('admin.change_admin_password'))

            current_user.set_password(new_password)
            current_user.updated_at = datetime.utcnow()
            try:
                db.session.commit()
                flash('Password changed successfully!', 'success')
                return redirect(url_for('admin.admin_profile'))
            except Exception as e:
                db.session.rollback()
                flash(f'Error changing password: {str(e)}', 'danger')

        return render_template('admin/change_password.html')
    except Exception as e:
        print(f"ERROR in change_admin_password: {str(e)}")
        traceback.print_exc()
        flash('Error changing password.', 'danger')
        return redirect(url_for('admin.admin_profile'))


@admin_bp.route('/settings', methods=['GET', 'POST'])
@admin_required
def admin_settings():
    try:
        if request.method == 'POST':
            flash('Settings updated successfully!', 'success')
            return redirect(url_for('admin.admin_settings'))

        settings = {
            'email_notifications':       True,
            'auto_approve_students':     False,
            'enable_proctoring':         True,
            'default_pass_percentage':   60,
            'allow_student_registration': True,
            'require_email_verification': False,
            'max_exam_duration':         180,
            'allow_exam_retake':         False
        }
        return render_template('admin/settings.html', settings=settings)
    except Exception as e:
        print(f"ERROR in admin_settings: {str(e)}")
        traceback.print_exc()
        flash('Error loading settings.', 'danger')
        return redirect(url_for('admin.dashboard'))


@admin_bp.route('/activity-log', methods=['GET'])
@admin_required
def activity_log():
    try:
        page           = request.args.get('page', 1, type=int)
        recent_users   = User.query.order_by(User.created_at.desc()).limit(10).all()
        recent_exams   = Exam.query.order_by(Exam.created_at.desc()).limit(10).all()
        recent_results = ExamResult.query.order_by(ExamResult.submitted_at.desc()).limit(10).all()

        activities = []
        for user in recent_users:
            activities.append({'type': 'user_created',
                               'description': f'User {user.full_name} ({user.role.name}) was created',
                               'timestamp': user.created_at, 'icon': 'fa-user-plus', 'color': 'primary'})
        for exam in recent_exams:
            activities.append({'type': 'exam_created',
                               'description': f'Exam "{exam.title}" was created',
                               'timestamp': exam.created_at, 'icon': 'fa-file-alt', 'color': 'info'})
        for result in recent_results:
            activities.append({'type': 'exam_submitted',
                               'description': f'{result.student_user.full_name} submitted {result.exam.title}',
                               'timestamp': result.submitted_at, 'icon': 'fa-check-circle', 'color': 'success'})

        activities.sort(key=lambda x: x['timestamp'], reverse=True)
        activities = activities[:30]

        return render_template('admin/activity_log.html', activities=activities, page=page)
    except Exception as e:
        print(f"ERROR in activity_log: {str(e)}")
        traceback.print_exc()
        flash('Error loading activity log.', 'danger')
        return redirect(url_for('admin.dashboard'))


# ══════════════════════════════════════════════════════════════════════════
# STUDENT MANAGEMENT
# ══════════════════════════════════════════════════════════════════════════

@admin_bp.route('/students', methods=['GET'])
@admin_required
def manage_students():
    try:
        page         = request.args.get('page', 1, type=int)
        search       = request.args.get('search', '', type=str)
        class_filter = request.args.get('class_id', None)

        if class_filter:
            try:
                class_filter = int(class_filter)
            except (ValueError, TypeError):
                class_filter = None

        query = Student.query.filter_by(is_active=True)
        if search:
            query = query.join(User).filter(db.or_(
                User.first_name.ilike(f'%{search}%'),
                User.last_name.ilike(f'%{search}%'),
                Student.admission_number.ilike(f'%{search}%')
            ))
        if class_filter:
            query = query.filter_by(class_id=class_filter)

        students = query.paginate(page=page, per_page=20, error_out=False)
        classes  = Class.query.filter_by(is_active=True).all()

        return render_template('admin/students.html', students=students, classes=classes,
                               search=search, class_filter=class_filter)
    except Exception as e:
        print(f"ERROR in manage_students: {str(e)}")
        traceback.print_exc()
        flash(f'Error loading students: {str(e)}', 'danger')
        return redirect(url_for('admin.dashboard'))


@admin_bp.app_template_filter('time_ago')
def time_ago(timestamp):
    if not timestamp:
        return "N/A"
    diff    = datetime.utcnow() - timestamp
    seconds = diff.total_seconds()
    if seconds < 60:      return "Just now"
    if seconds < 3600:    return f"{int(seconds // 60)} minutes ago"
    if seconds < 86400:   return f"{int(seconds // 3600)} hours ago"
    return f"{int(seconds // 86400)} days ago"


@admin_bp.route('/student/create', methods=['GET', 'POST'])
@admin_required
def create_student():
    if request.method == 'POST':
        username         = request.form.get('username', '').strip()
        email            = request.form.get('email', '').strip()
        first_name       = request.form.get('first_name', '').strip()
        last_name        = request.form.get('last_name', '').strip()
        admission_number = request.form.get('admission_number', '').strip()
        class_id         = request.form.get('class_id', type=int)
        gender           = request.form.get('gender', 'Other')
        password         = request.form.get('password', '')

        if not all([username, first_name, last_name, admission_number, password]):
            flash('Username, first name, last name, admission number, and password are required.', 'danger')
            return redirect(url_for('admin.create_student'))
        if User.query.filter_by(username=username).first():
            flash('Username already exists.', 'danger')
            return redirect(url_for('admin.create_student'))
        if Student.query.filter_by(admission_number=admission_number).first():
            flash('Admission number already exists.', 'danger')
            return redirect(url_for('admin.create_student'))
        if email and User.query.filter_by(email=email).first():
            flash('Email already exists.', 'danger')
            return redirect(url_for('admin.create_student'))

        try:
            role = Role.query.filter_by(name='Student').first()
            user = User(username=username, email=email or None,
                        first_name=first_name, last_name=last_name,
                        gender=gender, role_id=role.id)
            user.set_password(password)
            db.session.add(user)
            db.session.flush()

            student = Student(user_id=user.id, admission_number=admission_number,
                              class_id=class_id or None)
            db.session.add(student)
            db.session.commit()
            flash('Student created successfully.', 'success')
            return redirect(url_for('admin.manage_students'))
        except Exception as e:
            db.session.rollback()
            flash(f'Error creating student: {str(e)}', 'danger')

    classes = Class.query.filter_by(is_active=True).all()
    return render_template('admin/student_form.html', classes=classes, student=None)


@admin_bp.route('/student/<int:id>/edit', methods=['GET', 'POST'])
@admin_required
def edit_student(id):
    try:
        student = Student.query.get_or_404(id)
        if request.method == 'POST':
            student.user.first_name = request.form.get('first_name', '').strip()
            student.user.last_name  = request.form.get('last_name', '').strip()

            new_email = request.form.get('email', '').strip()
            if new_email:
                existing = User.query.filter_by(email=new_email).first()
                if existing and existing.id != student.user_id:
                    flash('Email already exists.', 'danger')
                    return redirect(url_for('admin.edit_student', id=id))
                student.user.email = new_email
            else:
                student.user.email = None

            student.user.gender = request.form.get('gender', 'Other')
            student.class_id    = request.form.get('class_id', type=int) or None

            new_username = request.form.get('username', '').strip()
            if new_username and new_username != student.user.username:
                existing = User.query.filter_by(username=new_username).first()
                if existing:
                    flash('Username already exists.', 'danger')
                    return redirect(url_for('admin.edit_student', id=id))
                student.user.username = new_username

            new_password = request.form.get('password', '').strip()
            if new_password:
                student.user.set_password(new_password)
                flash('Password updated successfully!', 'info')

            try:
                student.user.updated_at = datetime.utcnow()
                db.session.commit()
                flash('Student updated successfully.', 'success')
                return redirect(url_for('admin.manage_students'))
            except Exception as e:
                db.session.rollback()
                flash(f'Error updating student: {str(e)}', 'danger')

        classes = Class.query.filter_by(is_active=True).all()
        return render_template('admin/student_form.html', classes=classes, student=student)
    except Exception as e:
        print(f"ERROR in edit_student: {str(e)}")
        traceback.print_exc()
        flash(f'Error loading student: {str(e)}', 'danger')
        return redirect(url_for('admin.manage_students'))


@admin_bp.route('/student/<int:id>/delete', methods=['POST'])
@admin_required
def delete_student(id):
    student = Student.query.get_or_404(id)
    try:
        student.is_active      = False
        student.user.is_active = False
        db.session.commit()
        flash('Student deleted successfully.', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Error deleting student: {str(e)}', 'danger')
    return redirect(url_for('admin.manage_students'))


# ══════════════════════════════════════════════════════════════════════════
# TEACHER MANAGEMENT
# ══════════════════════════════════════════════════════════════════════════

@admin_bp.route('/teachers', methods=['GET'])
@admin_required
def manage_teachers():
    try:
        page           = request.args.get('page', 1, type=int)
        search         = request.args.get('search', '', type=str)
        subject_filter = request.args.get('subject', '', type=str)

        query = Teacher.query.filter_by(is_active=True)
        if search:
            query = query.join(User).filter(db.or_(
                User.first_name.ilike(f'%{search}%'),
                User.last_name.ilike(f'%{search}%'),
                Teacher.teacher_id.ilike(f'%{search}%')
            ))
        if subject_filter:
            query = query.filter_by(subject=subject_filter)

        teachers = query.paginate(page=page, per_page=20, error_out=False)
        subjects = sorted(list(set(t.subject for t in Teacher.query.filter_by(is_active=True).all() if t.subject)))

        return render_template('admin/teachers.html', teachers=teachers, subjects=subjects,
                               search=search, subject_filter=subject_filter)
    except Exception as e:
        print(f"ERROR in manage_teachers: {str(e)}")
        traceback.print_exc()
        flash(f'Error loading teachers: {str(e)}', 'danger')
        return redirect(url_for('admin.dashboard'))


@admin_bp.route('/teacher/create', methods=['GET', 'POST'])
@admin_required
def create_teacher():
    if request.method == 'POST':
        username       = request.form.get('username', '').strip()
        email          = request.form.get('email', '').strip()
        first_name     = request.form.get('first_name', '').strip()
        last_name      = request.form.get('last_name', '').strip()
        teacher_id     = request.form.get('teacher_id', '').strip()
        subject        = request.form.get('subject', '').strip()
        gender         = request.form.get('gender', 'Other')
        password       = request.form.get('password', '')
        qualification  = request.form.get('qualification', '').strip()
        specialization = request.form.get('specialization', '').strip()

        if not all([username, first_name, last_name, teacher_id, subject, password]):
            flash('Username, first name, last name, teacher ID, subject, and password are required.', 'danger')
            return redirect(url_for('admin.create_teacher'))
        if User.query.filter_by(username=username).first():
            flash('Username already exists.', 'danger')
            return redirect(url_for('admin.create_teacher'))
        if Teacher.query.filter_by(teacher_id=teacher_id).first():
            flash('Teacher ID already exists.', 'danger')
            return redirect(url_for('admin.create_teacher'))
        if email and User.query.filter_by(email=email).first():
            flash('Email already exists.', 'danger')
            return redirect(url_for('admin.create_teacher'))

        try:
            role = Role.query.filter_by(name='Teacher').first()
            user = User(username=username, email=email or None,
                        first_name=first_name, last_name=last_name,
                        gender=gender, role_id=role.id)
            user.set_password(password)
            db.session.add(user)
            db.session.flush()

            teacher = Teacher(user_id=user.id, teacher_id=teacher_id, subject=subject,
                              qualification=qualification or None,
                              specialization=specialization or None)
            db.session.add(teacher)
            db.session.commit()
            flash('Teacher created successfully.', 'success')
            return redirect(url_for('admin.manage_teachers'))
        except Exception as e:
            db.session.rollback()
            flash(f'Error creating teacher: {str(e)}', 'danger')

    return render_template('admin/teacher_form.html', teacher=None)


@admin_bp.route('/teacher/<int:id>/edit', methods=['GET', 'POST'])
@admin_required
def edit_teacher(id):
    try:
        teacher = Teacher.query.get_or_404(id)
        if request.method == 'POST':
            teacher.user.first_name = request.form.get('first_name', '').strip()
            teacher.user.last_name  = request.form.get('last_name', '').strip()

            new_email = request.form.get('email', '').strip()
            if new_email:
                existing = User.query.filter_by(email=new_email).first()
                if existing and existing.id != teacher.user_id:
                    flash('Email already exists.', 'danger')
                    return redirect(url_for('admin.edit_teacher', id=id))
                teacher.user.email = new_email
            else:
                teacher.user.email = None

            teacher.user.gender      = request.form.get('gender', 'Other')
            teacher.subject          = request.form.get('subject', '').strip()
            teacher.qualification    = request.form.get('qualification', '').strip()
            teacher.specialization   = request.form.get('specialization', '').strip()

            new_username = request.form.get('username', '').strip()
            if new_username and new_username != teacher.user.username:
                existing = User.query.filter_by(username=new_username).first()
                if existing:
                    flash('Username already exists.', 'danger')
                    return redirect(url_for('admin.edit_teacher', id=id))
                teacher.user.username = new_username

            new_password = request.form.get('password', '').strip()
            if new_password:
                teacher.user.set_password(new_password)
                flash('Password updated successfully!', 'info')

            try:
                teacher.user.updated_at = datetime.utcnow()
                db.session.commit()
                flash('Teacher updated successfully.', 'success')
                return redirect(url_for('admin.manage_teachers'))
            except Exception as e:
                db.session.rollback()
                flash(f'Error updating teacher: {str(e)}', 'danger')

        return render_template('admin/teacher_form.html', teacher=teacher)
    except Exception as e:
        print(f"ERROR in edit_teacher: {str(e)}")
        traceback.print_exc()
        flash(f'Error loading teacher: {str(e)}', 'danger')
        return redirect(url_for('admin.manage_teachers'))


@admin_bp.route('/teacher/<int:id>/delete', methods=['POST'])
@admin_required
def delete_teacher(id):
    teacher = Teacher.query.get_or_404(id)
    try:
        teacher.is_active      = False
        teacher.user.is_active = False
        db.session.commit()
        flash('Teacher deleted successfully.', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Error deleting teacher: {str(e)}', 'danger')
    return redirect(url_for('admin.manage_teachers'))


# ══════════════════════════════════════════════════════════════════════════
# USER MANAGEMENT
# ══════════════════════════════════════════════════════════════════════════

@admin_bp.route('/users')
@admin_required
def manage_users():
    try:
        page        = request.args.get('page', 1, type=int)
        search      = request.args.get('search', '', type=str)
        role_filter = request.args.get('role', '', type=str)

        query = User.query.filter_by(is_deleted=False)
        if search:
            query = query.filter(db.or_(
                User.username.ilike(f'%{search}%'),
                User.email.ilike(f'%{search}%'),
                User.first_name.ilike(f'%{search}%'),
                User.last_name.ilike(f'%{search}%')
            ))
        if role_filter:
            query = query.join(User.role).filter(Role.name == role_filter)

        users = query.order_by(User.created_at.desc()).paginate(page=page, per_page=20, error_out=False)
        return render_template('admin/manage_users.html', users=users,
                               search=search, role_filter=role_filter)
    except Exception as e:
        print(f"ERROR in manage_users: {str(e)}")
        traceback.print_exc()
        flash('Error loading users.', 'danger')
        return redirect(url_for('admin.dashboard'))


@admin_bp.route('/user/<int:user_id>/edit', methods=['GET', 'POST'])
@admin_required
def edit_user(user_id):
    try:
        user            = User.query.get_or_404(user_id)
        student_profile = Student.query.filter_by(user_id=user_id).first()
        teacher_profile = Teacher.query.filter_by(user_id=user_id).first()

        if request.method == 'POST':
            user.first_name = request.form.get('first_name')
            user.last_name  = request.form.get('last_name')
            user.email      = request.form.get('email')
            user.gender     = request.form.get('gender')

            new_username = request.form.get('username')
            if new_username != user.username:
                existing = User.query.filter_by(username=new_username).first()
                if existing:
                    flash('Username already exists.', 'danger')
                    return redirect(url_for('admin.edit_user', user_id=user_id))
                user.username = new_username

            new_password = request.form.get('new_password', '').strip()
            if new_password:
                user.set_password(new_password)
                flash(f'Password updated for {user.username}!', 'success')

            if student_profile:
                student_profile.contact_number  = request.form.get('contact_number')
                student_profile.address         = request.form.get('address')
                student_profile.guardian_name   = request.form.get('guardian_name')
                student_profile.guardian_contact = request.form.get('guardian_contact')

            if teacher_profile:
                teacher_profile.subject        = request.form.get('subject')
                teacher_profile.contact_number = request.form.get('contact_number')
                teacher_profile.address        = request.form.get('address')
                teacher_profile.qualification  = request.form.get('qualification')
                teacher_profile.specialization = request.form.get('specialization')

            user.updated_at = datetime.utcnow()
            db.session.commit()
            flash(f'User {user.username} updated successfully!', 'success')
            return redirect(url_for('admin.manage_users'))

        return render_template('admin/edit_user.html', user=user,
                               student_profile=student_profile,
                               teacher_profile=teacher_profile)
    except Exception as e:
        db.session.rollback()
        print(f"ERROR in edit_user: {str(e)}")
        traceback.print_exc()
        flash(f'Error updating user: {str(e)}', 'danger')
        return redirect(url_for('admin.manage_users'))


@admin_bp.route('/user/<int:user_id>/reset-password', methods=['POST'])
@admin_required
def reset_user_password(user_id):
    try:
        user         = User.query.get_or_404(user_id)
        new_password = request.form.get('new_password', '').strip()
        if not new_password:
            flash('Password cannot be empty.', 'danger')
            return redirect(url_for('admin.manage_users'))
        user.set_password(new_password)
        user.updated_at = datetime.utcnow()
        db.session.commit()
        flash(f'Password reset successfully for {user.username}!', 'success')
        return redirect(url_for('admin.manage_users'))
    except Exception as e:
        db.session.rollback()
        print(f"ERROR in reset_user_password: {str(e)}")
        flash('Error resetting password.', 'danger')
        return redirect(url_for('admin.manage_users'))


@admin_bp.route('/user/<int:user_id>/toggle-active', methods=['POST'])
@admin_required
def toggle_user_active(user_id):
    try:
        user = User.query.get_or_404(user_id)
        if user.id == current_user.id:
            flash('You cannot deactivate your own account.', 'warning')
            return redirect(url_for('admin.manage_users'))
        user.is_active  = not user.is_active
        user.updated_at = datetime.utcnow()
        db.session.commit()
        status = "activated" if user.is_active else "deactivated"
        flash(f'User {user.username} has been {status}!', 'success')
        return redirect(url_for('admin.manage_users'))
    except Exception as e:
        db.session.rollback()
        print(f"ERROR in toggle_user_active: {str(e)}")
        flash('Error updating user status.', 'danger')
        return redirect(url_for('admin.manage_users'))


@admin_bp.route('/user/<int:user_id>/delete', methods=['POST'])
@admin_required
def delete_user(user_id):
    try:
        user = User.query.get_or_404(user_id)
        if user.id == current_user.id:
            flash('You cannot delete your own account.', 'warning')
            return redirect(url_for('admin.manage_users'))
        user.is_deleted = True
        user.is_active  = False
        user.updated_at = datetime.utcnow()
        db.session.commit()
        flash(f'User {user.username} has been deleted!', 'success')
        return redirect(url_for('admin.manage_users'))
    except Exception as e:
        db.session.rollback()
        print(f"ERROR in delete_user: {str(e)}")
        flash('Error deleting user.', 'danger')
        return redirect(url_for('admin.manage_users'))


# ══════════════════════════════════════════════════════════════════════════
# BULK IMPORT
# ══════════════════════════════════════════════════════════════════════════

@admin_bp.route('/import/students', methods=['GET', 'POST'])
@admin_required
def import_students():
    if request.method == 'POST':
        if 'file' not in request.files or request.files['file'].filename == '':
            flash('No file selected.', 'danger')
            return redirect(url_for('admin.import_students'))
        try:
            filename, file_path = save_upload_file(request.files['file'], 'imports')
            results = import_students_from_csv(file_path, class_id=None)
            return render_template('admin/import_results.html', results=results, import_type='Students')
        except Exception as e:
            flash(f'Error importing file: {str(e)}', 'danger')
    return render_template('admin/import_students.html')


@admin_bp.route('/import/teachers', methods=['GET', 'POST'])
@admin_required
def import_teachers():
    if request.method == 'POST':
        if 'file' not in request.files or request.files['file'].filename == '':
            flash('No file selected.', 'danger')
            return redirect(url_for('admin.import_teachers'))
        try:
            filename, file_path = save_upload_file(request.files['file'], 'imports')
            results = import_teachers_from_csv(file_path)
            return render_template('admin/import_results.html', results=results, import_type='Teachers')
        except Exception as e:
            flash(f'Error importing file: {str(e)}', 'danger')
    return render_template('admin/import_teachers.html')


@admin_bp.route('/download/student-template')
@admin_required
def download_student_template():
    df     = generate_student_template()
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False)
    output.seek(0)
    return send_file(output,
                     mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                     as_attachment=True, download_name='student_template.xlsx')


@admin_bp.route('/download/teacher-template')
@admin_required
def download_teacher_template():
    df     = generate_teacher_template()
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False)
    output.seek(0)
    return send_file(output,
                     mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                     as_attachment=True, download_name='teacher_template.xlsx')


# ══════════════════════════════════════════════════════════════════════════
# EXAM MANAGEMENT
# ══════════════════════════════════════════════════════════════════════════

@admin_bp.route('/exams', methods=['GET'])
@admin_required
def manage_exams():
    try:
        page          = request.args.get('page', 1, type=int)
        search        = request.args.get('search', '', type=str)
        status_filter = request.args.get('status', '', type=str)

        query = Exam.query.filter_by(is_deleted=False)
        if search:
            query = query.filter(Exam.title.ilike(f'%{search}%'))
        if status_filter == 'published':
            query = query.filter_by(published=True)
        elif status_filter == 'unpublished':
            query = query.filter_by(published=False)

        exams = query.paginate(page=page, per_page=20, error_out=False)
        return render_template('admin/exams.html', exams=exams,
                               search=search, status_filter=status_filter)
    except Exception as e:
        print(f"ERROR in manage_exams: {str(e)}")
        traceback.print_exc()
        flash(f'Error loading exams: {str(e)}', 'danger')
        return redirect(url_for('admin.dashboard'))


@admin_bp.route('/exam/<int:id>/results', methods=['GET'])
@admin_required
def exam_results(id):
    """
    View exam results.
    FIX: Added explicit join(User) before the search filter so the User columns
         are available; without it the ilike filter on User columns caused an error.
    """
    try:
        exam    = Exam.query.get_or_404(id)
        page    = request.args.get('page', 1, type=int)
        search  = request.args.get('search', '', type=str)
        sort_by = request.args.get('sort', 'marks_obtained', type=str)

        # ── FIX: always join User so the search filter works ──
        query = ExamResult.query.filter_by(exam_id=id).join(
            User, ExamResult.student_id == User.id
        )

        if search:
            query = query.filter(db.or_(
                User.first_name.ilike(f'%{search}%'),
                User.last_name.ilike(f'%{search}%')
            ))

        if sort_by == 'marks_obtained':
            query = query.order_by(ExamResult.marks_obtained.desc())
        elif sort_by == 'percentage':
            query = query.order_by(ExamResult.percentage.desc())
        elif sort_by == 'submitted_at':
            query = query.order_by(ExamResult.submitted_at.desc())

        results     = query.paginate(page=page, per_page=20, error_out=False)
        all_results = ExamResult.query.filter_by(exam_id=id).all()

        if all_results:
            percentages = [r.percentage for r in all_results]
            avg_score   = sum(percentages) / len(percentages)
        else:
            avg_score = 0

        stats = {
            'total_attempts': len(all_results),
            'avg_score':      round(avg_score, 2),
            'pass_count':     sum(1 for r in all_results if r.is_passed),
            'fail_count':     sum(1 for r in all_results if not r.is_passed)
        }

        return render_template('admin/exam_results.html', exam=exam, results=results,
                               stats=stats, search=search, sort_by=sort_by)
    except Exception as e:
        print(f"ERROR in exam_results: {str(e)}")
        traceback.print_exc()
        flash(f'Error loading exam results: {str(e)}', 'danger')
        return redirect(url_for('admin.manage_exams'))


@admin_bp.route('/exam/<int:result_id>/result-details')
@admin_required
def result_details_view(result_id):
    """
    View detailed result.
    FIX 1: unanswered count — was `len(answers) - correct - wrong` which goes
            negative when an answer is neither correct nor wrong (unselected).
            Now counted directly.
    FIX 2: correct_answer_text — used getattr(question,'correct_answer') which
            is None for MCQ. Now uses get_correct_answer_text() helper.
    FIX 3: student_answer_text — now uses get_student_answer_text() helper.
    FIX 4: question/option order — uses get_ordered_questions() which respects
            the per-session randomised order stored in ExamSession.question_order.
    """
    try:
        result       = ExamResult.query.get_or_404(result_id)
        exam_session = ExamSession.query.get(result.exam_session_id)

        answers = StudentAnswer.query.filter_by(
            exam_session_id=result.exam_session_id,
            student_id=result.student_id
        ).all()

        answer_map = {a.question_id: a for a in answers}

        # ── FIX: correct counts ──
        correct    = sum(1 for a in answers if a.is_correct)
        wrong      = sum(1 for a in answers
                         if not a.is_correct and (a.selected_option_id or a.theory_answer))
        unanswered = sum(1 for a in answers
                         if not a.selected_option_id and not a.theory_answer)

        statistics = {'correct': correct, 'wrong': wrong, 'unanswered': unanswered}

        # ── FIX: get questions in the order the student saw them ──
        questions = get_ordered_questions(result.exam_id, exam_session)

        question_results = []
        for question in questions:
            answer = answer_map.get(question.id)

            if answer and answer.is_correct:
                status = 'correct'
            elif answer and (answer.selected_option_id or answer.theory_answer):
                status = 'wrong'
            else:
                status = 'unanswered'

            # ── FIX: proper text for both MCQ and theory ──
            student_answer_text = get_student_answer_text(answer, question)
            correct_answer_text = get_correct_answer_text(question)

            question_results.append({
                'question':           question,
                'answer':             answer,
                'status':             status,
                'student_answer_text': student_answer_text,
                'correct_answer_text': correct_answer_text,
                'marks_obtained':     (answer.marks_obtained or 0) if answer else 0
            })

        return render_template('admin/result_details.html',
                               result=result,
                               answers=answers,
                               statistics=statistics,
                               question_results=question_results)
    except Exception as e:
        print(f"ERROR in result_details: {str(e)}")
        traceback.print_exc()
        flash(f'Error loading result details: {str(e)}', 'danger')
        return redirect(url_for('admin.manage_exams'))


# ══════════════════════════════════════════════════════════════════════════
# PDF RESULT GENERATION
# ══════════════════════════════════════════════════════════════════════════

@admin_bp.route('/result/<int:result_id>/generate-pdf')
@admin_required
def generate_result_pdf(result_id):
    """
    Generate a detailed PDF for one exam result.
    FIX: Questions and options are now rendered in the order the student
         actually saw them (respecting ExamSession.question_order and
         ExamSession.option_order JSON fields).
    """
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib import colors
        from reportlab.lib.units import mm
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
            HRFlowable, KeepTogether, PageBreak
        )
        import re

        # ── helpers ──────────────────────────────────────────────────────────

        def strip_latex(text):
            if not text:
                return ""
            text = str(text)
            text = re.sub(r'\$\$(.+?)\$\$', r'\1', text, flags=re.DOTALL)
            text = re.sub(r'\$(.+?)\$',     r'\1', text, flags=re.DOTALL)
            text = re.sub(r'\\\((.+?)\\\)', r'\1', text, flags=re.DOTALL)
            text = re.sub(r'\\\[(.+?)\\\]', r'\1', text, flags=re.DOTALL)
            text = re.sub(r'\\frac\{([^}]+)\}\{([^}]+)\}', r'(\1/\2)', text)
            text = re.sub(r'\\sqrt\{([^}]+)\}',            r'√(\1)',    text)
            text = re.sub(r'\\text\{([^}]+)\}',            r'\1',       text)
            text = re.sub(r'\\mathrm\{([^}]+)\}',          r'\1',       text)
            text = re.sub(r'\\left|\\right',               '',          text)
            text = text.replace(r'\times', '×').replace(r'\div', '÷')
            text = text.replace(r'\leq', '≤').replace(r'\geq', '≥')
            text = text.replace(r'\neq', '≠').replace(r'\approx', '≈')
            text = text.replace(r'\infty', '∞').replace(r'\pi', 'π')
            text = text.replace(r'\alpha', 'α').replace(r'\beta', 'β')
            text = text.replace(r'\theta', 'θ').replace(r'\lambda', 'λ')
            text = text.replace(r'\mu', 'μ').replace(r'\sigma', 'σ')
            text = text.replace(r'\Delta', 'Δ').replace(r'\Sigma', 'Σ')
            text = re.sub(r'_\{([^}]+)\}',  r'_\1',  text)
            text = re.sub(r'\^\{([^}]+)\}', r'^\1',  text)
            text = re.sub(r'\\[a-zA-Z]+',   '',      text)
            text = re.sub(r'\{|\}',          '',      text)
            return text.strip()

        def para(text, style):
            safe = strip_latex(str(text))
            safe = safe.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            return Paragraph(safe, style)

        # ── data ─────────────────────────────────────────────────────────────

        result       = ExamResult.query.get_or_404(result_id)
        exam         = result.exam
        student      = result.student_user
        exam_session = ExamSession.query.get(result.exam_session_id)

        answers    = StudentAnswer.query.filter_by(
            exam_session_id=result.exam_session_id,
            student_id=result.student_id
        ).all()
        answer_map = {a.question_id: a for a in answers}

        # ── FIX: use session-aware ordering ──
        questions = get_ordered_questions(exam.id, exam_session)

        n_total      = len(questions)
        n_correct    = sum(1 for a in answers if a.is_correct)
        n_wrong      = sum(1 for a in answers
                           if not a.is_correct and (a.selected_option_id or a.theory_answer))
        n_unanswered = sum(1 for a in answers
                           if not a.selected_option_id and not a.theory_answer)

        # ── colour palette ────────────────────────────────────────────────────
        C_DARK      = colors.HexColor('#0d1117')
        C_GOLD      = colors.HexColor('#c9973a')
        C_EMERALD   = colors.HexColor('#0d7a5f')
        C_EMERALD_L = colors.HexColor('#e6f4f0')
        C_CRIMSON   = colors.HexColor('#c0392b')
        C_CRIMSON_L = colors.HexColor('#fbeae8')
        C_AMBER     = colors.HexColor('#b7791f')
        C_AMBER_L   = colors.HexColor('#fef9ec')
        C_BLUE      = colors.HexColor('#1e40af')
        C_BLUE_L    = colors.HexColor('#eff3ff')
        C_GREY_L    = colors.HexColor('#f2ede6')
        C_GREY_B    = colors.HexColor('#e8e2d9')
        C_WHITE     = colors.white
        C_OPT_N     = colors.HexColor('#f9fafb')

        # ── styles ────────────────────────────────────────────────────────────
        def S(name, **kw):
            d = dict(fontName='Helvetica', fontSize=10, leading=14,
                     textColor=C_DARK, spaceAfter=0, spaceBefore=0)
            d.update(kw)
            return ParagraphStyle(name, **d)

        sTitle   = S('Title',   fontName='Helvetica-Bold', fontSize=18,
                     textColor=C_WHITE, alignment=TA_CENTER, leading=24)
        sSub     = S('Sub',     fontName='Helvetica', fontSize=9,
                     textColor=colors.HexColor('#8b949e'), alignment=TA_CENTER)
        sH1      = S('H1',      fontName='Helvetica-Bold', fontSize=13, textColor=C_DARK, spaceAfter=4)
        sH2      = S('H2',      fontName='Helvetica-Bold', fontSize=11, textColor=C_DARK)
        sBody    = S('Body',    fontSize=10, leading=14, textColor=colors.HexColor('#3d4151'))
        sBodyB   = S('BodyB',   fontName='Helvetica-Bold', fontSize=10, textColor=C_DARK)
        sSmall   = S('Small',   fontSize=8.5, textColor=colors.HexColor('#6e7491'))
        sCenter  = S('Center',  fontSize=10, alignment=TA_CENTER, textColor=C_DARK)
        sCenterB = S('CenterB', fontName='Helvetica-Bold', fontSize=10,
                     alignment=TA_CENTER, textColor=C_DARK)
        sQText   = S('QText',   fontSize=9.5, leading=14, textColor=C_DARK)

        # ── document ──────────────────────────────────────────────────────────
        buf    = io.BytesIO()
        PAGE_W, PAGE_H = A4
        MARGIN = 18 * mm

        def add_footer(canvas, doc):
            canvas.saveState()
            canvas.setFont('Helvetica', 8)
            canvas.setFillColor(colors.HexColor('#6e7491'))
            ts = datetime.now().strftime('%Y-%m-%d %H:%M')
            canvas.drawString(MARGIN, 10 * mm, f'Generated: {ts}')
            canvas.drawRightString(PAGE_W - MARGIN, 10 * mm, f'Page {doc.page}')
            canvas.setStrokeColor(C_GREY_B)
            canvas.setLineWidth(0.5)
            canvas.line(MARGIN, 13 * mm, PAGE_W - MARGIN, 13 * mm)
            canvas.restoreState()

        doc   = SimpleDocTemplate(buf, pagesize=A4,
                                  leftMargin=MARGIN, rightMargin=MARGIN,
                                  topMargin=MARGIN,  bottomMargin=22 * mm,
                                  title=f"Result — {student.full_name}",
                                  author="E-Exam Portal")
        story = []
        TW    = PAGE_W - 2 * MARGIN

        # ── header ────────────────────────────────────────────────────────────
        hdr_tbl = Table([[para('EXAMINATION RESULT REPORT', sTitle)]], colWidths=[TW])
        hdr_tbl.setStyle(TableStyle([
            ('BACKGROUND',   (0,0), (-1,-1), C_DARK),
            ('TOPPADDING',   (0,0), (-1,-1), 14),
            ('BOTTOMPADDING',(0,0), (-1,-1), 14),
            ('LEFTPADDING',  (0,0), (-1,-1), 10),
            ('RIGHTPADDING', (0,0), (-1,-1), 10),
        ]))
        story.append(hdr_tbl)
        story.append(Spacer(1, 3*mm))
        story.append(para(f"Generated: {datetime.now().strftime('%A, %d %B %Y  %H:%M:%S')}", sSub))
        story.append(Spacer(1, 6*mm))

        # ── info panel ────────────────────────────────────────────────────────
        def kv_row(label, value, vs=sBody):
            return [para(label, sSmall), para(str(value), vs)]

        def make_info_table(rows):
            t = Table(rows, colWidths=[TW*0.38, TW*0.62])
            t.setStyle(TableStyle([
                ('BACKGROUND',   (0,0), (-1,0),  C_GREY_L),
                ('FONTNAME',     (0,0), (-1,0),  'Helvetica-Bold'),
                ('SPAN',         (0,0), (-1,0)),
                ('TOPPADDING',   (0,0), (-1,-1), 5),
                ('BOTTOMPADDING',(0,0), (-1,-1), 5),
                ('LEFTPADDING',  (0,0), (-1,-1), 8),
                ('RIGHTPADDING', (0,0), (-1,-1), 8),
                ('LINEBELOW',    (0,0), (-1,-1), 0.4, C_GREY_B),
                ('VALIGN',       (0,0), (-1,-1), 'TOP'),
            ]))
            return t

        student_rows = [
            [para('STUDENT INFORMATION', sH2), ''],
            kv_row('Full Name',  student.full_name,  sBodyB),
            kv_row('Email',      student.email or '—'),
            kv_row('Submitted',  result.submitted_at.strftime('%d %b %Y  %H:%M:%S')),
        ]
        exam_rows = [
            [para('EXAM INFORMATION', sH2), ''],
            kv_row('Exam Title',  exam.title,         sBodyB),
            kv_row('Total Marks', result.total_marks),
            kv_row('Pass Marks',  result.pass_marks),
        ]

        info_panel = Table([[make_info_table(student_rows), make_info_table(exam_rows)]],
                           colWidths=[TW*0.50 - 3, TW*0.50 - 3])
        info_panel.setStyle(TableStyle([
            ('LEFTPADDING',  (0,0), (-1,-1), 0),
            ('RIGHTPADDING', (0,0), (-1,-1), 0),
            ('TOPPADDING',   (0,0), (-1,-1), 0),
            ('BOTTOMPADDING',(0,0), (-1,-1), 0),
            ('ALIGN',        (0,0), (-1,-1), 'LEFT'),
            ('VALIGN',       (0,0), (-1,-1), 'TOP'),
        ]))
        story.append(info_panel)
        story.append(Spacer(1, 5*mm))

        # ── score panel ───────────────────────────────────────────────────────
        pct_color = C_EMERALD if result.is_passed else C_CRIMSON
        pct_bg    = C_EMERALD_L if result.is_passed else C_CRIMSON_L
        verdict   = 'PASSED ✓' if result.is_passed else 'FAILED ✗'

        sScoreCol = S('ScoreCol', fontName='Helvetica-Bold', fontSize=28,
                      alignment=TA_CENTER, leading=34, textColor=pct_color)
        sVerdictS = S('VerdictS', fontName='Helvetica-Bold', fontSize=11,
                      alignment=TA_CENTER, textColor=pct_color)

        score_data = [[
            Table([[para(f"{result.percentage}%", sScoreCol)],
                   [para(verdict, sVerdictS)],
                   [para(f"{result.marks_obtained} / {result.total_marks} marks", sCenterB)]],
                  colWidths=[TW * 0.32]),
            Table([[para('Grade', sSmall),     para(result.grade or 'N/A', sCenterB)],
                   [para('Pass Mark', sSmall), para(str(result.pass_marks), sCenter)],
                   [para('Total Qs', sSmall),  para(str(n_total), sCenter)]],
                  colWidths=[TW*0.18, TW*0.18]),
        ]]
        score_tbl = Table(score_data, colWidths=[TW*0.35, TW*0.65])
        score_tbl.setStyle(TableStyle([
            ('BACKGROUND',   (0,0), (0,-1), pct_bg),
            ('BACKGROUND',   (1,0), (1,-1), C_GREY_L),
            ('ALIGN',        (0,0), (-1,-1), 'CENTER'),
            ('VALIGN',       (0,0), (-1,-1), 'MIDDLE'),
            ('TOPPADDING',   (0,0), (-1,-1), 10),
            ('BOTTOMPADDING',(0,0), (-1,-1), 10),
            ('LEFTPADDING',  (0,0), (-1,-1), 8),
            ('RIGHTPADDING', (0,0), (-1,-1), 8),
            ('BOX',          (0,0), (-1,-1), 0.8, C_GREY_B),
            ('INNERGRID',    (0,0), (-1,-1), 0.4, C_GREY_B),
        ]))
        story.append(score_tbl)
        story.append(Spacer(1, 4*mm))

        # ── stats bar ─────────────────────────────────────────────────────────
        def stat_cell(num, label, bg, fg):
            sN = S(f'SN{label}', fontName='Helvetica-Bold', fontSize=16,
                   alignment=TA_CENTER, textColor=fg)
            sL = S(f'SL{label}', fontSize=8, alignment=TA_CENTER, textColor=fg)
            t  = Table([[para(str(num), sN)], [para(label, sL)]], colWidths=[TW/3 - 4])
            t.setStyle(TableStyle([
                ('BACKGROUND',   (0,0), (-1,-1), bg),
                ('TOPPADDING',   (0,0), (-1,-1), 8),
                ('BOTTOMPADDING',(0,0), (-1,-1), 8),
            ]))
            return t

        stat_row = Table([[
            stat_cell(n_correct,   'Correct Answers', C_EMERALD_L, C_EMERALD),
            stat_cell(n_wrong,     'Wrong Answers',   C_CRIMSON_L, C_CRIMSON),
            stat_cell(n_unanswered,'Unanswered',      C_AMBER_L,   C_AMBER),
        ]], colWidths=[TW/3 - 2]*3, hAlign='CENTER')
        stat_row.setStyle(TableStyle([
            ('LEFTPADDING',  (0,0), (-1,-1), 2),
            ('RIGHTPADDING', (0,0), (-1,-1), 2),
            ('TOPPADDING',   (0,0), (-1,-1), 0),
            ('BOTTOMPADDING',(0,0), (-1,-1), 0),
        ]))
        story.append(stat_row)
        story.append(Spacer(1, 7*mm))

        # ── section heading ───────────────────────────────────────────────────
        story.append(HRFlowable(width=TW, thickness=1, color=C_GREY_B))
        story.append(Spacer(1, 3*mm))
        story.append(para('Question-by-Question Analysis', sH1))
        story.append(Spacer(1, 3*mm))

        # ── legend ────────────────────────────────────────────────────────────
        def legend_chip(label, bg, fg):
            s = S(f'leg{label}', fontSize=7.5, alignment=TA_CENTER,
                  textColor=fg, fontName='Helvetica-Bold')
            t = Table([[para(label, s)]], colWidths=[30*mm])
            t.setStyle(TableStyle([
                ('BACKGROUND',   (0,0), (-1,-1), bg),
                ('TOPPADDING',   (0,0), (-1,-1), 3),
                ('BOTTOMPADDING',(0,0), (-1,-1), 3),
            ]))
            return t

        legend = Table([[
            legend_chip('✓ Correct Answer',         C_EMERALD_L, C_EMERALD),
            legend_chip('✗ Student Wrong Selection', C_CRIMSON_L, C_CRIMSON),
            legend_chip('○ Other Options',           C_OPT_N,     colors.HexColor('#6e7491')),
        ]], colWidths=[34*mm]*3, hAlign='LEFT')
        legend.setStyle(TableStyle([
            ('LEFTPADDING',  (0,0), (-1,-1), 2),
            ('RIGHTPADDING', (0,0), (-1,-1), 2),
            ('TOPPADDING',   (0,0), (-1,-1), 0),
            ('BOTTOMPADDING',(0,0), (-1,-1), 0),
        ]))
        story.append(legend)
        story.append(Spacer(1, 4*mm))

        # ── question table header ─────────────────────────────────────────────
        COL_NUM  = 8  * mm
        COL_Q    = TW * 0.52
        COL_ANS  = TW * 0.18
        COL_CORR = TW * 0.18
        COL_MRK  = TW - COL_NUM - COL_Q - COL_ANS - COL_CORR

        hdr = Table([[
            para('#',                  sCenterB),
            para('Question & Options', sBodyB),
            para('Your Answer',        sCenterB),
            para('Correct',            sCenterB),
            para('Marks',              sCenterB),
        ]], colWidths=[COL_NUM, COL_Q, COL_ANS, COL_CORR, COL_MRK])
        hdr.setStyle(TableStyle([
            ('BACKGROUND',   (0,0), (-1,-1), C_DARK),
            ('TEXTCOLOR',    (0,0), (-1,-1), C_WHITE),
            ('FONTNAME',     (0,0), (-1,-1), 'Helvetica-Bold'),
            ('FONTSIZE',     (0,0), (-1,-1), 8.5),
            ('TOPPADDING',   (0,0), (-1,-1), 6),
            ('BOTTOMPADDING',(0,0), (-1,-1), 6),
            ('LEFTPADDING',  (0,0), (-1,-1), 5),
            ('ALIGN',        (0,0), (-1,-1), 'CENTER'),
            ('ALIGN',        (1,0), (1,0),   'LEFT'),
        ]))
        story.append(hdr)

        LABELS = ['A', 'B', 'C', 'D', 'E', 'F']

        # ── per-question rows ─────────────────────────────────────────────────
        for idx, question in enumerate(questions, 1):
            answer = answer_map.get(question.id)

            if answer and answer.is_correct:
                row_bg = C_EMERALD_L; status = 'correct'
            elif answer and (answer.selected_option_id or answer.theory_answer):
                row_bg = C_CRIMSON_L; status = 'wrong'
            else:
                row_bg = C_AMBER_L;   status = 'skip'

            row_stripe = C_WHITE if idx % 2 == 0 else colors.HexColor('#fdfcfa')

            opt_rows     = []
            correct_text = '—'
            student_text = '—'

            if question.question_type in ('mcq', 'true_false'):
                # ── FIX: use session-aware option order ──
                options = get_ordered_options(question, exam_session)

                for i, opt in enumerate(options):
                    lbl  = opt.option_label or (LABELS[i] if i < len(LABELS) else str(i+1))
                    otxt = strip_latex(opt.option_text or '')
                    sel  = answer and answer.selected_option_id == opt.id
                    corr = bool(opt.is_correct)

                    if corr:
                        correct_text = f"{lbl}. {otxt}"
                        ob = C_EMERALD_L; ofg = C_EMERALD; marker = '✓'
                    elif sel and not corr:
                        student_text = f"{lbl}. {otxt}"
                        ob = C_CRIMSON_L; ofg = C_CRIMSON; marker = '✗'
                    else:
                        ob = C_OPT_N; ofg = colors.HexColor('#6e7491'); marker = '○'

                    sOL = S(f'OL{idx}{i}', fontName='Helvetica-Bold', fontSize=8, textColor=ofg)
                    sOT = S(f'OT{idx}{i}', fontSize=8, leading=11,
                            textColor=C_DARK if (corr or sel) else colors.HexColor('#4b5563'))

                    row_t = Table([[para(f'{marker} {lbl}.', sOL), para(otxt, sOT)]],
                                  colWidths=[10*mm, COL_Q - 12*mm])
                    row_t.setStyle(TableStyle([
                        ('BACKGROUND',   (0,0), (-1,-1), ob),
                        ('TOPPADDING',   (0,0), (-1,-1), 3),
                        ('BOTTOMPADDING',(0,0), (-1,-1), 3),
                        ('LEFTPADDING',  (0,0), (-1,-1), 4),
                        ('RIGHTPADDING', (0,0), (-1,-1), 4),
                        ('VALIGN',       (0,0), (-1,-1), 'TOP'),
                    ]))
                    opt_rows.append(row_t)

                # Resolve student_text if not yet set by wrong-option path
                if answer and answer.selected_option_id and student_text == '—':
                    sel_opt = next((o for o in options if o.id == answer.selected_option_id), None)
                    if sel_opt:
                        student_text = f"{sel_opt.option_label}. {strip_latex(sel_opt.option_text or '')}"

            else:
                # Theory
                student_text = strip_latex(answer.theory_answer if answer and answer.theory_answer else '—')
                correct_text = strip_latex(getattr(question, 'correct_answer', None) or '—')

            # Build question cell
            q_cell_items = [para(strip_latex(question.question_text or ''), sQText),
                            Spacer(1, 2*mm)]
            for ot in opt_rows:
                q_cell_items.append(ot)
                q_cell_items.append(Spacer(1, 1))

            q_inner = Table([[item] for item in q_cell_items], colWidths=[COL_Q])
            q_inner.setStyle(TableStyle([
                ('LEFTPADDING',  (0,0), (-1,-1), 0),
                ('RIGHTPADDING', (0,0), (-1,-1), 0),
                ('TOPPADDING',   (0,0), (-1,-1), 0),
                ('BOTTOMPADDING',(0,0), (-1,-1), 0),
            ]))

            mk_fg = C_EMERALD if status == 'correct' else C_CRIMSON if status == 'wrong' else C_AMBER
            sMK   = S(f'MK{idx}',  fontName='Helvetica-Bold', fontSize=10,
                      alignment=TA_CENTER, textColor=mk_fg)
            sMKs  = S(f'MKS{idx}', fontSize=7.5, alignment=TA_CENTER,
                      textColor=colors.HexColor('#6e7491'))
            mk_cell = Table([[para(str(answer.marks_obtained if answer else 0), sMK)],
                             [para(f'/ {question.marks}', sMKs)]],
                            colWidths=[COL_MRK])
            mk_cell.setStyle(TableStyle([
                ('TOPPADDING',   (0,0), (-1,-1), 2),
                ('BOTTOMPADDING',(0,0), (-1,-1), 2),
                ('LEFTPADDING',  (0,0), (-1,-1), 2),
                ('RIGHTPADDING', (0,0), (-1,-1), 2),
            ]))

            sNum  = S(f'Num{idx}',  fontName='Helvetica-Bold', fontSize=9,
                      alignment=TA_CENTER, textColor=mk_fg)
            sAns  = S(f'Ans{idx}',  fontSize=8, leading=12, textColor=mk_fg)
            sCorr = S(f'Corr{idx}', fontSize=8, leading=12,
                      fontName='Helvetica-Bold', textColor=C_EMERALD)

            q_row = Table([[para(str(idx), sNum), q_inner,
                            para(student_text, sAns), para(correct_text, sCorr), mk_cell]],
                          colWidths=[COL_NUM, COL_Q, COL_ANS, COL_CORR, COL_MRK])
            q_row.setStyle(TableStyle([
                ('BACKGROUND',   (0,0), (-1,-1), row_stripe),
                ('TOPPADDING',   (0,0), (-1,-1), 7),
                ('BOTTOMPADDING',(0,0), (-1,-1), 7),
                ('LEFTPADDING',  (0,0), (-1,-1), 5),
                ('RIGHTPADDING', (0,0), (-1,-1), 5),
                ('VALIGN',       (0,0), (-1,-1), 'TOP'),
                ('LINEBELOW',    (0,0), (-1,-1), 0.5, C_GREY_B),
                ('LINEBEFORE',   (0,0), (0,-1),  3,   mk_fg),
            ]))
            story.append(KeepTogether(q_row))

        # ── totals row ────────────────────────────────────────────────────────
        story.append(Spacer(1, 3*mm))
        total_obtained = sum((a.marks_obtained or 0) for a in answers)
        tot_row = Table([[
            para('', sCenter), para('TOTAL', sCenterB),
            para('', sCenter), para('', sCenter),
            para(f'{total_obtained}\n/ {result.total_marks}', sCenterB),
        ]], colWidths=[COL_NUM, COL_Q, COL_ANS, COL_CORR, COL_MRK])
        tot_row.setStyle(TableStyle([
            ('BACKGROUND',   (0,0), (-1,-1), C_DARK),
            ('TEXTCOLOR',    (0,0), (-1,-1), C_WHITE),
            ('TOPPADDING',   (0,0), (-1,-1), 8),
            ('BOTTOMPADDING',(0,0), (-1,-1), 8),
            ('LEFTPADDING',  (0,0), (-1,-1), 5),
            ('ALIGN',        (0,0), (-1,-1), 'CENTER'),
        ]))
        story.append(tot_row)

        doc.build(story, onFirstPage=add_footer, onLaterPages=add_footer)
        pdf_bytes = buf.getvalue()
        buf.close()

        fname    = (f"Result_{student.full_name.replace(' ','_')}"
                    f"_{exam.title.replace(' ','_')}.pdf")
        response = make_response(pdf_bytes)
        response.headers.set('Content-Type', 'application/pdf')
        response.headers.set('Content-Disposition', 'attachment', filename=fname)
        return response

    except Exception as e:
        traceback.print_exc()
        flash(f'Error generating PDF: {str(e)}', 'danger')
        return redirect(url_for('admin.result_details_view', result_id=result_id))


# ══════════════════════════════════════════════════════════════════════════
# PROCTORING REPORTS
# ══════════════════════════════════════════════════════════════════════════

@admin_bp.route('/proctoring/reports')
@login_required
@admin_required
def proctoring_reports():
    exam_id        = request.args.get('exam_id', type=int)
    student_id     = request.args.get('student_id', type=int)
    violation_type = request.args.get('violation_type')
    date_from      = request.args.get('date_from')
    date_to        = request.args.get('date_to')
    page           = request.args.get('page', 1, type=int)

    query = db.session.query(ProctoringLog).join(
        ExamSession, ProctoringLog.exam_session_id == ExamSession.id
    ).join(User, ExamSession.student_id == User.id).join(Exam, ExamSession.exam_id == Exam.id)

    if exam_id:
        query = query.filter(ExamSession.exam_id == exam_id)
    if student_id:
        query = query.filter(ExamSession.student_id == student_id)
    if violation_type:
        field = 'violation_type' if hasattr(ProctoringLog, 'violation_type') else 'event_type'
        query = query.filter(getattr(ProctoringLog, field) == violation_type)
    if date_from:
        query = query.filter(ProctoringLog.timestamp >= datetime.strptime(date_from, '%Y-%m-%d'))
    if date_to:
        query = query.filter(ProctoringLog.timestamp <= datetime.strptime(date_to, '%Y-%m-%d'))

    query              = query.order_by(ProctoringLog.timestamp.desc())
    logs_pagination    = query.paginate(page=page, per_page=50, error_out=False)

    for log in logs_pagination.items:
        es = ExamSession.query.get(log.exam_session_id)
        if es:
            log.student_user = User.query.get(es.student_id)
            log.exam         = Exam.query.get(es.exam_id)
        else:
            log.student_user = log.exam = None

    stats      = calculate_proctoring_stats()
    exams      = Exam.query.filter_by(published=True).order_by(Exam.title).all()
    students   = db.session.query(User).join(Student, User.id == Student.user_id).order_by(
        User.first_name, User.last_name).all()
    violation_types = [
        'face_not_visible', 'multiple_faces', 'tab_switch', 'fullscreen_exit',
        'copy_attempt', 'paste_attempt', 'excessive_violations',
        'camera_access_denied', 'dev_tools_attempt'
    ]

    return render_template('admin/proctoring_reports.html',
                           logs=logs_pagination, stats=stats, exams=exams,
                           students=students, violation_types=violation_types,
                           filters={'exam_id': exam_id, 'student_id': student_id,
                                    'violation_type': violation_type,
                                    'date_from': date_from, 'date_to': date_to})


def calculate_proctoring_stats():
    stats = {}
    use_vt = hasattr(ProctoringLog, 'violation_type')
    field  = 'violation_type' if use_vt else 'event_type'

    for key, value in [('face_not_visible', 'face_not_visible'),
                       ('multiple_faces',   'multiple_faces'),
                       ('tab_switch',       'tab_switch'),
                       ('fullscreen_exit',  'fullscreen_exit')]:
        stats[key] = ProctoringLog.query.filter(
            getattr(ProctoringLog, field) == value
        ).count()

    if use_vt:
        stats['total_violations'] = ProctoringLog.query.filter(
            ProctoringLog.violation_type.isnot(None)).count()
        if hasattr(ProctoringLog, 'severity'):
            stats['suspicious_activity'] = ProctoringLog.query.filter_by(severity='high').count()
        else:
            stats['suspicious_activity'] = ProctoringLog.query.filter(
                ProctoringLog.violation_type.in_(
                    ['multiple_faces', 'excessive_violations', 'dev_tools_attempt']
                )).count()
    else:
        stats['total_violations']   = ProctoringLog.query.count()
        stats['suspicious_activity'] = ProctoringLog.query.filter(
            ProctoringLog.event_type.in_(
                ['multiple_faces', 'excessive_violations', 'dev_tools_attempt']
            )).count()

    return stats


@admin_bp.route('/proctoring/student/<int:student_id>/exam/<int:exam_id>')
@login_required
@admin_required
def student_proctoring_detail(student_id, exam_id):
    student_user    = User.query.get_or_404(student_id)
    student_profile = Student.query.filter_by(user_id=student_id).first_or_404()
    exam            = Exam.query.get_or_404(exam_id)
    exam_session    = ExamSession.query.filter_by(
        student_id=student_id, exam_id=exam_id).first_or_404()
    logs = ProctoringLog.query.filter_by(
        exam_session_id=exam_session.id
    ).order_by(ProctoringLog.timestamp.asc()).all()

    violation_summary = {
        'tab_switches':   exam_session.tab_switches   or 0,
        'copy_attempts':  exam_session.copy_attempts  or 0,
        'paste_attempts': exam_session.paste_attempts or 0,
    }
    if hasattr(exam_session, 'face_violations'):
        violation_summary['face_violations']  = exam_session.face_violations  or 0
    if hasattr(exam_session, 'fullscreen_exits'):
        violation_summary['fullscreen_exits'] = exam_session.fullscreen_exits or 0

    total_violations = sum(violation_summary.values())

    return render_template('admin/student_proctoring_detail.html',
                           student=student_user, student_profile=student_profile,
                           exam=exam, exam_session=exam_session, logs=logs,
                           violation_summary=violation_summary,
                           total_violations=total_violations)


@admin_bp.route('/proctoring/export/<int:exam_id>')
@login_required
@admin_required
def export_proctoring_report(exam_id):
    try:
        exam = Exam.query.get_or_404(exam_id)
        logs = db.session.query(ProctoringLog).join(
            ExamSession, ProctoringLog.exam_session_id == ExamSession.id
        ).filter(ExamSession.exam_id == exam_id).join(
            User, ExamSession.student_id == User.id
        ).order_by(ProctoringLog.timestamp.asc()).all()

        if not logs:
            flash('No proctoring data found for this exam.', 'warning')
            return redirect(url_for('admin.proctoring_reports'))

        wb = Workbook()
        ws = wb.active
        ws.title = "Proctoring Report"
        for col, w in zip(['A','B','C','D','E','F'], [20,25,20,12,40,15]):
            ws.column_dimensions[col].width = w

        hfill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
        hfont = Font(bold=True, color="FFFFFF", size=12)
        halign = Alignment(horizontal="center", vertical="center")
        bdr    = Border(left=Side(style='thin'), right=Side(style='thin'),
                        top=Side(style='thin'),  bottom=Side(style='thin'))

        ws.merge_cells('A1:F1')
        ws['A1'].value     = f"Proctoring Report: {exam.title}"
        ws['A1'].font      = Font(bold=True, size=14, color="4472C4")
        ws['A1'].alignment = Alignment(horizontal="center", vertical="center")
        ws.merge_cells('A2:F2')
        ws['A2'].value     = f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        ws['A2'].font      = Font(size=10, italic=True)
        ws['A2'].alignment = Alignment(horizontal="center")

        for col, hdr in enumerate(['Timestamp','Student Name','Violation Type',
                                   'Severity','Details','Session ID'], start=1):
            c = ws.cell(row=4, column=col)
            c.value = hdr; c.fill = hfill; c.font = hfont
            c.alignment = halign; c.border = bdr

        row = 5
        for log in logs:
            es           = ExamSession.query.get(log.exam_session_id)
            student_name = User.query.get(es.student_id).full_name if es else "Unknown"
            vtype        = (getattr(log, 'violation_type', None) or
                            getattr(log, 'event_type', 'event'))
            severity     = getattr(log, 'severity', '-') or '-'
            details      = getattr(log, 'details', 'N/A') or 'N/A'
            if len(details) > 100:
                details = details[:100] + '...'

            ws.cell(row=row, column=1).value = log.timestamp.strftime('%Y-%m-%d %H:%M:%S')
            ws.cell(row=row, column=2).value = student_name
            ws.cell(row=row, column=3).value = vtype.replace('_',' ').title() if vtype else 'Event'
            ws.cell(row=row, column=4).value = severity.upper() if severity else '-'
            ws.cell(row=row, column=5).value = details
            ws.cell(row=row, column=6).value = str(log.exam_session_id)

            for col in range(1, 7):
                c = ws.cell(row=row, column=col)
                c.border = bdr
                c.alignment = Alignment(vertical="center", wrap_text=True)

            sc = ws.cell(row=row, column=4)
            sl = severity.lower() if severity else ''
            if   sl == 'high':   sc.fill = PatternFill(start_color="FFC7CE",end_color="FFC7CE",fill_type="solid"); sc.font = Font(color="9C0006",bold=True)
            elif sl == 'medium': sc.fill = PatternFill(start_color="FFEB9C",end_color="FFEB9C",fill_type="solid"); sc.font = Font(color="9C6500",bold=True)
            elif sl == 'low':    sc.fill = PatternFill(start_color="C6EFCE",end_color="C6EFCE",fill_type="solid"); sc.font = Font(color="006100")
            row += 1

        # Summary
        row += 2
        ws.merge_cells(f'A{row}:F{row}')
        ws[f'A{row}'].value     = "Summary Statistics"
        ws[f'A{row}'].font      = Font(bold=True, size=12, color="4472C4")
        ws[f'A{row}'].alignment = Alignment(horizontal="center")
        row += 1

        vcounts = {}
        scounts = {'high': 0, 'medium': 0, 'low': 0}
        ustudents = set()
        for log in logs:
            vt = getattr(log, 'violation_type', None) or getattr(log, 'event_type', None)
            if vt: vcounts[vt] = vcounts.get(vt, 0) + 1
            sv = getattr(log, 'severity', None)
            if sv: scounts[sv.lower()] = scounts.get(sv.lower(), 0) + 1
            es = ExamSession.query.get(log.exam_session_id)
            if es: ustudents.add(es.student_id)

        for label, value in [
            ('Total Violations:', len(logs)),
            ('Students with Violations:', len(ustudents)),
            ('High Severity:', scounts['high']),
            ('Medium Severity:', scounts['medium']),
            ('Low Severity:', scounts['low']),
        ]:
            ws.cell(row=row, column=1).value = label
            ws.cell(row=row, column=1).font  = Font(bold=True)
            ws.cell(row=row, column=2).value = value
            row += 1

        row += 1
        ws.cell(row=row, column=1).value = "Violation Breakdown:"
        ws.cell(row=row, column=1).font  = Font(bold=True)
        row += 1
        for vt, cnt in sorted(vcounts.items(), key=lambda x: x[1], reverse=True):
            ws.cell(row=row, column=1).value = vt.replace('_',' ').title()
            ws.cell(row=row, column=2).value = cnt
            row += 1

        temp_dir = tempfile.gettempdir()
        filename = f"proctoring_report_{exam.id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
        filepath = os.path.join(temp_dir, filename)
        wb.save(filepath)

        return send_file(filepath, as_attachment=True, download_name=filename,
                         mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')

    except Exception as e:
        print(f"ERROR in export_proctoring_report: {str(e)}")
        traceback.print_exc()
        flash(f'Error exporting report: {str(e)}', 'danger')
        return redirect(url_for('admin.proctoring_reports'))


# ══════════════════════════════════════════════════════════════════════════
# STUDENT REPORTS
# ══════════════════════════════════════════════════════════════════════════

@admin_bp.route('/reports/students', methods=['GET', 'POST'])
@admin_required
def student_reports():
    try:
        if request.method == 'POST':
            class_id    = request.form.get('class_id', type=int)
            exam_id     = request.form.get('exam_id', type=int)
            report_type = request.form.get('report_type', 'summary')

            if not class_id and not exam_id:
                flash('Please select at least a class or an exam.', 'warning')
                return redirect(url_for('admin.student_reports'))

            dispatch = {
                'summary':     generate_summary_report,
                'detailed':    generate_detailed_report,
                'performance': generate_performance_report,
            }
            fn = dispatch.get(report_type)
            if fn:
                return fn(class_id, exam_id)
            flash('Invalid report type selected.', 'danger')
            return redirect(url_for('admin.student_reports'))

        classes = Class.query.filter_by(is_active=True).order_by(Class.name).all()
        exams   = Exam.query.filter_by(is_deleted=False).order_by(Exam.created_at.desc()).all()
        return render_template('admin/student_reports.html', classes=classes, exams=exams)
    except Exception as e:
        print(f"ERROR in student_reports: {str(e)}")
        traceback.print_exc()
        flash(f'Error loading reports page: {str(e)}', 'danger')
        return redirect(url_for('admin.dashboard'))


def generate_summary_report(class_id, exam_id):
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        report_data = get_report_data(class_id, exam_id)
        if not report_data['students']:
            flash('No data found for the selected filters.', 'warning')
            return redirect(url_for('admin.student_reports'))

        doc   = Document()
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(11)

        title     = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run(report_data['title'])
        title_run.font.size = Pt(18); title_run.font.bold = True

        class_name = report_data['class'].name if report_data['class'] else 'All Classes'
        exam_title = report_data['exam'].title if report_data['exam'] else 'All Exams'
        meta       = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta.add_run(f"Class: {class_name}\nExam: {exam_title}\n")
        meta.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}").italic = True
        doc.add_paragraph()

        doc.add_heading('Summary Statistics', level=2)
        stats_tbl = doc.add_table(rows=len(report_data['statistics']) + 1, cols=2)
        stats_tbl.style = 'Light Grid Accent 1'
        for cell in stats_tbl.rows[0].cells:
            cell.text = ['Metric', 'Value'][stats_tbl.rows[0].cells.index(cell)]
            for p in cell.paragraphs:
                for r in p.runs: r.font.bold = True
        for idx, (k, v) in enumerate(report_data['statistics'].items(), start=1):
            row = stats_tbl.rows[idx].cells
            row[0].text = format_key(k); row[1].text = format_value(v)
        doc.add_paragraph()

        doc.add_heading('Student Results', level=2)
        tbl = doc.add_table(rows=len(report_data['students']) + 1, cols=5)
        tbl.style = 'Light Grid Accent 1'
        for cell, h in zip(tbl.rows[0].cells,
                           ['Admission No.','Student Name','Class','Marks/Percentage','Status']):
            cell.text = h
            for p in cell.paragraphs:
                for r in p.runs: r.font.bold = True
        for idx, sd in enumerate(report_data['students'], start=1):
            row = tbl.rows[idx].cells
            row[0].text = sd['student'].admission_number
            row[1].text = sd['student'].user.full_name
            row[2].text = sd['class_name']
            row[3].text = (f"{sd['marks']} ({sd['percentage']})" if sd['marks'] != 'N/A'
                           else sd['percentage'])
            row[4].text = sd['status']

        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename  = (f"Student_Report_{class_name.replace(' ','_')}"
                     f"_{exam_title.replace(' ','_')}_{timestamp}.docx")
        doc.save(f'/mnt/user-data/outputs/{filename}')
        flash(f'Report "{filename}" generated successfully!', 'success')
        return redirect(url_for('admin.student_reports'))
    except Exception as e:
        print(f"ERROR in generate_summary_report: {str(e)}")
        traceback.print_exc()
        flash(f'Error generating report: {str(e)}', 'danger')
        return redirect(url_for('admin.student_reports'))


def generate_detailed_report(class_id, exam_id):
    """
    FIX: answer.selected_option → correct attribute is selected_option_id; look up
         option text via QuestionOption.query.get(). Also fixed correct_answer lookup
         for MCQ (was reading nonexistent question.correct_answer field).
    FIX: question.order used for row numbering; question.question_number removed.
    """
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        report_data = get_report_data(class_id, exam_id)
        if not report_data['students']:
            flash('No data found for the selected filters.', 'warning')
            return redirect(url_for('admin.student_reports'))
        if not exam_id:
            flash('Detailed reports require a specific exam.', 'warning')
            return redirect(url_for('admin.student_reports'))

        doc  = Document()
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tr = title.add_run('Detailed Student Performance Report')
        tr.font.size = Pt(18); tr.font.bold = True

        exam = report_data['exam']
        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta.add_run(f"Exam: {exam.title}\n"
                     f"Class: {report_data['class'].name if report_data['class'] else 'All'}\n")
        meta.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y')}").italic = True
        doc.add_paragraph()

        # Questions in default exam order (no session here — bulk report)
        questions = get_ordered_questions(exam_id)

        for sd in report_data['students']:
            student = sd['student']
            doc.add_heading(f"{student.user.full_name} ({student.admission_number})", level=2)

            result = ExamResult.query.filter_by(
                student_id=student.user_id, exam_id=exam_id).first()

            if not result:
                doc.add_paragraph("Did not attempt this exam.")
                doc.add_page_break()
                continue

            # Per-student session for randomised order
            student_session = ExamSession.query.filter_by(
                student_id=student.user_id, exam_id=exam_id).first()
            student_questions = get_ordered_questions(exam_id, student_session)

            summary = doc.add_paragraph()
            summary.add_run(
                f"Total Score: {result.marks_obtained}/{result.total_marks} "
                f"({result.percentage:.2f}%)\n"
            ).bold = True
            summary.add_run(
                f"Status: {'Passed' if result.is_passed else 'Failed'}\n"
                f"Submitted: {result.submitted_at.strftime('%B %d, %Y at %I:%M %p')}"
            )

            doc.add_heading('Question Analysis', level=3)

            answers = StudentAnswer.query.filter_by(
                exam_session_id=result.exam_session_id,
                student_id=student.user_id
            ).all()
            answer_dict = {a.question_id: a for a in answers}

            q_tbl = doc.add_table(rows=len(student_questions) + 1, cols=4)
            q_tbl.style = 'Light Grid Accent 1'
            for cell, h in zip(q_tbl.rows[0].cells,
                               ['Q#', 'Student Answer', 'Correct Answer', 'Result']):
                cell.text = h
                for p in cell.paragraphs:
                    for r in p.runs: r.font.bold = True

            for idx, question in enumerate(student_questions, start=1):
                row    = q_tbl.rows[idx].cells
                row[0].text = str(idx)   # show display index, not DB order field

                answer = answer_dict.get(question.id)

                # ── FIX: get student answer text properly ──
                row[1].text = get_student_answer_text(answer, question)
                # ── FIX: get correct answer for MCQ from options ──
                row[2].text = get_correct_answer_text(question)
                row[3].text = ('✓ Correct' if answer and answer.is_correct
                               else '✗ Wrong' if answer else '✗ Unanswered')

            doc.add_page_break()

        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename  = f"Detailed_Student_Report_{timestamp}.docx"
        doc.save(f'/mnt/user-data/outputs/{filename}')
        flash(f'Detailed report "{filename}" generated successfully!', 'success')
        return redirect(url_for('admin.student_reports'))
    except Exception as e:
        print(f"ERROR in generate_detailed_report: {str(e)}")
        traceback.print_exc()
        flash(f'Error generating detailed report: {str(e)}', 'danger')
        return redirect(url_for('admin.student_reports'))


def generate_performance_report(class_id, exam_id):
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        report_data = get_report_data(class_id, exam_id)
        if not report_data['students']:
            flash('No data found for the selected filters.', 'warning')
            return redirect(url_for('admin.student_reports'))

        doc  = Document()
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tr = title.add_run('Performance Analysis Report')
        tr.font.size = Pt(18); tr.font.bold = True

        class_name = report_data['class'].name if report_data['class'] else 'All Classes'
        exam_title = report_data['exam'].title if report_data['exam'] else 'All Exams'
        meta       = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta.add_run(f"Class: {class_name}\nExam: {exam_title}\n")
        meta.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y')}").italic = True
        doc.add_paragraph()

        doc.add_heading('Overall Statistics', level=2)
        stats_tbl = doc.add_table(rows=len(report_data['statistics']) + 1, cols=2)
        stats_tbl.style = 'Light Grid Accent 1'
        for cell, h in zip(stats_tbl.rows[0].cells, ['Metric', 'Value']):
            cell.text = h
            for p in cell.paragraphs:
                for r in p.runs: r.font.bold = True
        for idx, (k, v) in enumerate(report_data['statistics'].items(), start=1):
            row = stats_tbl.rows[idx].cells
            row[0].text = format_key(k); row[1].text = format_value(v)
        doc.add_paragraph()

        doc.add_heading('Performance Distribution', level=2)
        if exam_id:
            grades      = {'A (90-100%)': 0, 'B (80-89%)': 0, 'C (70-79%)': 0,
                           'D (60-69%)': 0, 'F (0-59%)': 0}
            all_results = ExamResult.query.filter_by(exam_id=exam_id).all()
            if class_id:
                # FIX: use student.user_id (not student_id key)
                student_ids = [sd['student'].user_id for sd in report_data['students']]
                all_results = [r for r in all_results if r.student_id in student_ids]
            for r in all_results:
                if   r.percentage >= 90: grades['A (90-100%)'] += 1
                elif r.percentage >= 80: grades['B (80-89%)']  += 1
                elif r.percentage >= 70: grades['C (70-79%)']  += 1
                elif r.percentage >= 60: grades['D (60-69%)']  += 1
                else:                    grades['F (0-59%)']   += 1

            g_tbl = doc.add_table(rows=len(grades) + 1, cols=2)
            g_tbl.style = 'Light Grid Accent 1'
            for cell, h in zip(g_tbl.rows[0].cells, ['Grade', 'Number of Students']):
                cell.text = h
                for p in cell.paragraphs:
                    for r in p.runs: r.font.bold = True
            for idx, (grade, count) in enumerate(grades.items(), start=1):
                row = g_tbl.rows[idx].cells
                row[0].text = grade; row[1].text = str(count)

        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename  = f"Performance_Report_{timestamp}.docx"
        doc.save(f'/mnt/user-data/outputs/{filename}')
        flash(f'Performance report "{filename}" generated successfully!', 'success')
        return redirect(url_for('admin.student_reports'))
    except Exception as e:
        print(f"ERROR in generate_performance_report: {str(e)}")
        traceback.print_exc()
        flash(f'Error generating performance report: {str(e)}', 'danger')
        return redirect(url_for('admin.student_reports'))


def get_report_data(class_id, exam_id):
    data = {'class': None, 'exam': None, 'students': [], 'statistics': {}, 'title': ''}
    if class_id: data['class'] = Class.query.get(class_id)
    if exam_id:  data['exam']  = Exam.query.get(exam_id)

    if data['class'] and data['exam']:
        data['title'] = f"{data['class'].name} - {data['exam'].title} Student Report"
    elif data['class']:
        data['title'] = f"{data['class'].name} - Overall Student Report"
    elif data['exam']:
        data['title'] = f"{data['exam'].title} - Student Performance Report"
    else:
        data['title'] = "Student Report"

    query = Student.query.filter_by(is_active=True)
    if class_id: query = query.filter_by(class_id=class_id)
    students = query.all()

    for student in students:
        class_name = 'No Class'
        if student.class_id:
            sc = Class.query.get(student.class_id)
            if sc: class_name = sc.name

        sd = {'student': student, 'class_name': class_name,
              'marks': 'N/A', 'percentage': 'N/A', 'status': 'Not Attempted'}

        if exam_id:
            result = ExamResult.query.filter_by(
                student_id=student.user_id, exam_id=exam_id).first()
            if result:
                sd['marks']      = f"{result.marks_obtained}/{result.total_marks}"
                sd['percentage'] = f"{result.percentage:.2f}%"
                sd['status']     = 'Passed' if result.is_passed else 'Failed'
        else:
            results = ExamResult.query.filter_by(student_id=student.user_id).all()
            if results:
                avg = sum(r.percentage for r in results) / len(results)
                sd['percentage'] = f"{avg:.2f}%"
                sd['marks']      = f"{len(results)} exams"
                sd['status']     = f"{sum(1 for r in results if r.is_passed)}/{len(results)} passed"

        data['students'].append(sd)

    if exam_id:
        all_results = ExamResult.query.filter_by(exam_id=exam_id).all()
        if class_id:
            sids        = [s.user_id for s in students]
            all_results = [r for r in all_results if r.student_id in sids]
        if all_results:
            pcts = [r.percentage for r in all_results]
            data['statistics'] = {
                'Total Students':   len(students),
                'Total Attempts':   len(all_results),
                'Average Score':    f"{sum(pcts)/len(pcts):.2f}%",
                'Highest Score':    f"{max(pcts):.2f}%",
                'Lowest Score':     f"{min(pcts):.2f}%",
                'Students Passed':  sum(1 for r in all_results if r.is_passed),
                'Students Failed':  sum(1 for r in all_results if not r.is_passed),
                'Pass Rate':        f"{sum(1 for r in all_results if r.is_passed)/len(all_results)*100:.2f}%"
            }
        else:
            data['statistics'] = {'Total Students': len(students), 'Total Attempts': 0,
                                  'Average Score': '0%', 'Pass Rate': '0%'}
    else:
        data['statistics'] = {
            'Total Students': len(students),
            'Total Classes':  1 if class_id else Class.query.filter_by(is_active=True).count()
        }

    return data


def format_key(key):   return key.replace('_', ' ').title()
def format_value(val): return f"{val:.2f}" if isinstance(val, float) else str(val)


# ══════════════════════════════════════════════════════════════════════════
# MISC API ENDPOINTS
# ══════════════════════════════════════════════════════════════════════════

@admin_bp.route('/api/stats/students-count')
@admin_required
def api_students_count():
    try:
        count = Student.query.filter_by(is_active=True).count()
        return jsonify({'count': count})
    except Exception as e:
        return jsonify({'count': 0, 'error': str(e)}), 500


# ══════════════════════════════════════════════════════════════════════════
# DELETE RESULT FOR RETAKE
# ══════════════════════════════════════════════════════════════════════════

@admin_bp.route('/result/<int:result_id>/delete-for-retake', methods=['POST'])
@login_required
@admin_required
def delete_result_for_retake(result_id):
    try:
        print(f"\n=== DELETE RESULT FOR RETAKE ===")
        result     = ExamResult.query.get_or_404(result_id)
        exam       = result.exam
        student    = result.student_user
        session_id = result.exam_session_id

        StudentAnswer.query.filter_by(exam_session_id=session_id).delete()
        ProctoringLog.query.filter_by(exam_session_id=session_id).delete()

        es = ExamSession.query.get(session_id)
        if es: db.session.delete(es)
        db.session.delete(result)
        db.session.commit()

        print(f"✅ Deleted result {result_id} — {student.full_name} can retake {exam.title}")
        flash(f'Result deleted. {student.full_name} can now retake "{exam.title}".', 'success')
        return redirect(url_for('admin.exam_results', id=exam.id))
    except Exception as e:
        db.session.rollback()
        traceback.print_exc()
        flash(f'Error deleting result: {str(e)}', 'danger')
        try:
            return redirect(url_for('admin.exam_results', id=result.exam.id))
        except Exception:
            return redirect(url_for('admin.manage_exams'))


@admin_bp.route('/exam/<int:exam_id>/bulk-delete-for-retake', methods=['POST'])
@login_required
@admin_required
def bulk_delete_for_retake(exam_id):
    try:
        result_ids    = request.form.getlist('result_ids[]')
        if not result_ids:
            flash('No results selected.', 'warning')
            return redirect(url_for('admin.exam_results', id=exam_id))

        deleted_count = 0
        student_names = []
        for rid in result_ids:
            try:
                result = ExamResult.query.get(int(rid))
                if not result: continue
                sid = result.exam_session_id
                student_names.append(result.student_user.full_name)
                StudentAnswer.query.filter_by(exam_session_id=sid).delete()
                ProctoringLog.query.filter_by(exam_session_id=sid).delete()
                es = ExamSession.query.get(sid)
                if es: db.session.delete(es)
                db.session.delete(result)
                deleted_count += 1
            except Exception as e:
                print(f"Error deleting result {rid}: {str(e)}")
        db.session.commit()

        if deleted_count:
            flash(f'Deleted {deleted_count} result(s). Students can now retake the exam.', 'success')
        else:
            flash('No results were deleted.', 'warning')
        return redirect(url_for('admin.exam_results', id=exam_id))
    except Exception as e:
        db.session.rollback()
        traceback.print_exc()
        flash(f'Error during bulk deletion: {str(e)}', 'danger')
        return redirect(url_for('admin.exam_results', id=exam_id))


# ══════════════════════════════════════════════════════════════════════════
# ANALYTICS (simple dashboard)
# ══════════════════════════════════════════════════════════════════════════

@admin_bp.route('/analytics')
@admin_required
def analytics():
    try:
        all_results = ExamResult.query.all()

        class_performance = []
        for cls in Class.query.all():
            exams    = Exam.query.filter_by(class_id=cls.id).all()
            exam_ids = [e.id for e in exams]
            if exam_ids:
                results = ExamResult.query.filter(ExamResult.exam_id.in_(exam_ids)).all()
                if results:
                    avg = sum(r.percentage for r in results) / len(results)
                    class_performance.append((cls.name, avg, len(results)))

        student_scores = defaultdict(list)
        for r in all_results:
            student_scores[r.student_id].append(r.percentage)

        top_performers = []
        for sid, scores in student_scores.items():
            user = User.query.get(sid)
            if user:
                top_performers.append((user.first_name, user.last_name,
                                       sum(scores)/len(scores)))
        top_performers.sort(key=lambda x: x[2], reverse=True)
        top_performers = top_performers[:10]

        exam_stats = []
        for exam in Exam.query.all():
            results = ExamResult.query.filter_by(exam_id=exam.id).all()
            if results:
                pcts = [r.percentage for r in results]
                exam_stats.append((exam.title, len(results),
                                   sum(pcts)/len(pcts), max(pcts), min(pcts)))

        return render_template('admin/analytics.html',
                               class_performance=class_performance,
                               top_performers=top_performers,
                               exam_stats=exam_stats)
    except Exception as e:
        print(f"ERROR in analytics: {str(e)}")
        traceback.print_exc()
        flash(f'Error loading analytics: {str(e)}', 'danger')
        return redirect(url_for('admin.dashboard'))


@admin_bp.route('/api/analytics/chart-data')
@admin_required
def analytics_chart_data():
    try:
        chart_type = request.args.get('type', 'performance')
        if chart_type == 'performance':
            class_data = []
            for cls in Class.query.all():
                exams    = Exam.query.filter_by(class_id=cls.id).all()
                exam_ids = [e.id for e in exams]
                if exam_ids:
                    results = ExamResult.query.filter(ExamResult.exam_id.in_(exam_ids)).all()
                    if results:
                        avg = sum(r.percentage for r in results) / len(results)
                        class_data.append({'name': cls.name, 'score': avg})
            return jsonify({'labels': [d['name'] for d in class_data],
                            'data':   [round(d['score'], 2) for d in class_data]})
        return jsonify({'error': 'Invalid chart type'}), 400
    except Exception as e:
        print(f"ERROR in analytics_chart_data: {str(e)}")
        return jsonify({'error': str(e)}), 500


# ══════════════════════════════════════════════════════════════════════════
# PROCTORING VIOLATION LOG (API)
# FIX: was inside api_students_count scope due to wrong indentation —
#      now properly defined at module level.
# ══════════════════════════════════════════════════════════════════════════

@admin_bp.route('/exam/proctoring/log', methods=['POST'])
@login_required
def log_proctoring_violation():
    data = request.get_json() or {}

    session = ExamSession.query.filter_by(
        student_id=current_user.id,
        status='started'
    ).first()

    if not session:
        return jsonify({'status': 'error', 'message': 'No active exam session'}), 400

    violation_type = data.get('violation_type') or 'unknown'
    severity       = data.get('severity')       or 'medium'
    screenshot     = data.get('screenshot_path')

    log = ProctoringLog(
        exam_id          = session.exam_id,
        student_id       = current_user.id,
        exam_session_id  = session.id,
        violation_type   = violation_type,
        severity         = severity,
        timestamp        = datetime.utcnow(),
        screenshot_path  = screenshot
    )
    db.session.add(log)
    session.webcam_captures = (session.webcam_captures or 0) + 1
    db.session.commit()

    return jsonify({'status': 'ok'})